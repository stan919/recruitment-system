import os
import requests
from plantuml import PlantUML
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

pl = PlantUML(url='http://www.plantuml.com/plantuml/png/')

diagrams = {
    'usecase': '''@startuml
left to right direction
skinparam packageStyle rectangle

actor "求职者 Student" as S
actor "系统管理员 Admin" as A
actor "企业人事 HR" as C

rectangle "职引未来平台" {
  usecase "搜索过滤公开职位" as UC1
  usecase "宏观市场洞察查看" as UC2
  usecase "简历云端管理解析" as UC3
  usecase "投递及溯源追踪" as UC4
  usecase "即时通讯联络(Chat)" as UC5
  usecase "发布与维保职位" as UC6
  usecase "简历申请表审核" as UC7
  usecase "发起面谈沟通邀约" as UC8
  usecase "企业资质全局管控" as UC9
  usecase "全局用户及系统治理" as UC10
}

S --> UC1
S --> UC2
S --> UC3
S --> UC4
S --> UC5

C --> UC6
C --> UC7
C --> UC8

A --> UC9
A --> UC10
@enduml''',
    
    'architecture': '''@startuml
skinparam componentStyle rectangle
node "视图表现层 Views / Frontend" {
  component "Jinja2 & Bootstrap UI" as UI
  component "DOM原生层 Fetch API" as JS
  component "ECharts 大盘组件" as ECH
}

node "业务路由层 Controllers" {
  component "鉴权及拦截网" as AUTH
  component "Restful API(业务/通讯)" as API
  component "可视化静态计算引擎" as VIS
}

node "持久模型层 Models / Data" {
  component "SQLAlchemy 实体映射" as ORM
  database "MySQL 8.0 关系库" as DB
  component "Pandas 清洗任务" as DATA
}

UI --> AUTH : 跨路由拦截校验
JS --> API : JSON Payload 交互
ECH --> VIS : 聚合调用
AUTH --> ORM 
API --> ORM
VIS --> DATA : 组维与切片聚合
DATA --> DB : 离线数仓载入
ORM --> DB : ORM 持久化与追踪
@enduml''',

    'flowchart': '''@startuml
start
:用户点击“投递岗位”;
if (资料库判断:已完善个人简历?) then (否)
  :抛出异常并阻断通信动作;
  :重定向跳转至[个人简历配置中心];
else (是)
  :激活前置挂钩：验证岗位活跃度与容限阻断判定;
  :向 MySQL Application 锁入快照防撤回留底;
  :企业侧 HR 回台触发轮询更新收到游离端通告;
  if (实施第一道简历初步筛审流程) then (筛选出局/淘汰)
    :指令改写行状态点为 "拒绝 (Refused)";
    :退落至学生申请列表更新展示红标反馈;
  else (资质评估通过)
    :表状态升级调整至进阶标定为 "邀约 (Invited)";
    :系统解锁双方隔离聊天区边界;
    :正式步入私域动态即时通信(聊天Chat室)阶段;
end if
endif
stop
@enduml''',
    
    'class_diagram': '''@startuml
class User {
  +int id
  +string email
  +string password_hash
  +string role
  +boolean verify_password(pwd)
}
class Profile_Student {
  +int user_id
  +string real_name
  +string university
  +int graduation_year
}
class Resume {
  +int id
  +int student_id
  +string file_path
  +string parsed_content
}
class Company {
  +int id
  +int owner_id
  +string company_name
  +boolean is_verified
}
class Job {
  +int id
  +int company_id
  +string title
  +int salary_min
}

User "1" *-- "1" Profile_Student : "携带扩展类挂起集"
Profile_Student "1" *-- "0..*" Resume : "负责上行解析分发文件"
User "1" *-- "0..1" Company : "依角色判定专属特权域"
Company "1" *-- "0..*" Job : "发布及管控管理标的岗位"
@enduml''',
    
    'sequence_diagram': '''@startuml
participant "用户UI前端" as S
participant "后端 API网关" as API
participant "数据库底库" as DB
participant "企业端UI控制面板" as HR

S -> API : POST /apply指令挂载主键请求位
activate API
API -> API : 反构签署凭证，鉴阅与防重放逻辑
API -> DB : 溯回确认核验冗余及有效性上限
activate DB
DB --> API : 记录表合法回证无误放行
deactivate DB
API -> DB : 实施十字绑靠法装填 Application挂载快照操作
activate DB
DB --> API : 操作库事务落盘完成返回写入流
deactivate DB
API --> S : 返回正常200 交互绿标信通确认
deactivate API

autonumber
HR -> API : 发送常驻轮询 GET /applications
activate API
API -> DB : 并表映射挂列当前列下待审申请
activate DB
DB --> API : 获取主副文数据切片对象群
deactivate DB
API --> HR : 发配给视图页渲染列表列帧阵
deactivate API
HR -> API : 回发决定性触发指令 PUT 更改阶段态位邀约
activate API
API -> DB : 确认覆写表栏节点解围双信向静域
API --> HR : 对前端弹窗释放聊天跳点权限路由
deactivate API
@enduml''',
    
    'er_diagram': '''@startuml
entity "用户架构系 (USER)" as U {
  * id : INT <<PK>>
  --
  email : VARCHAR
  hashed_pwd : VARCHAR
  role : ENUM
}
entity "求职者拓扑表 (STUDENT_PROFILE)" as SP {
  * id : INT <<PK>>
  --
  user_id : INT <<FK>>
  real_name : VARCHAR
}
entity "法人实体制证系 (COMPANY)" as C {
  * id : INT <<PK>>
  --
  owner_id : INT <<FK>>
  company_name : VARCHAR
}
entity "发布集岗位池 (JOB)" as J {
  * id : INT <<PK>>
  --
  company_id : INT <<FK>>
  title : VARCHAR
  salary : INT
}
entity "枢纽流转链路单(APPLICATION)" as A {
  * id : INT <<PK>>
  --
  job_id : INT <<FK>>
  resume_id : INT <<FK>>
  status : VARCHAR
}

U ||..o| SP : "单个基底扩展对应一套简历组包"
U ||..o| C : "管理系限制强行映射法代表企业主域"
C ||..o{ J : "法主体具备广撒推流岗位池子项"
SP ||..o{ A : "跨库组合多份简历定向靶向投递追踪单"
J ||..o{ A : "独立节点接受接纳无数下流靶投单并绑定快查"
@enduml'''
}

os.makedirs('static/img', exist_ok=True)
images_paths = {}

print("Downloading generated diagrams from PlantUML API...")
for name, code in diagrams.items():
    try:
        raw_png = pl.processes(code)
        img_path = f"static/img/{name}.png"
        with open(img_path, 'wb') as f:
            f.write(raw_png)
        images_paths[name] = img_path
        print(f"  Saved {img_path}")
    except Exception as e:
        print(f"Error fetching {name}: {e}")

print("Updating README.md with diagram references...")
with open('README.md', 'r', encoding='utf-8') as f:
    readme_text = f.read()

if "![顶层用例模型]" not in readme_text:
    readme_text = readme_text.replace("### 1.5 初始顶层用例模型", 
    f"### 1.5 初始顶层用例模型\n\n![顶层用例模型]({images_paths.get('usecase', '')})\n\n")

if "![系统架构设计图]" not in readme_text:
    readme_text = readme_text.replace("### 3.1 架构设计决策与技术选型依据", 
    f"### 3.1 架构设计决策与技术选型依据\n\n#### 系统架构设计\n![系统架构设计图]({images_paths.get('architecture', '')})\n\n")

if "![系统流程图]" not in readme_text:
    readme_text = readme_text.replace("### 4.1 问题域核心业务编码", 
    f"### 4.1 问题域核心业务编码\n\n#### 系统核心投递与沟通流程图\n![系统流程图]({images_paths.get('flowchart', '')})\n\n")

if "![数据模型ER图]" not in readme_text:
    readme_text = readme_text.replace("### 3.2 实体关系与数据库模型架构设计", 
    f"### 3.2 实体关系与数据库模型架构设计\n\n#### 数据库 ER 模型图\n![数据模型ER图]({images_paths.get('er_diagram', '')})\n\n#### 业务领域静态模型（核心实体类图）\n![业务领域静态模型]({images_paths.get('class_diagram', '')})\n\n")

if "![业务涉及动态模型]" not in readme_text:
    readme_text = readme_text.replace("### 3.3 核心拓展模块与可行性设计凭证", 
    f"### 3.3 核心拓展模块与可行性设计凭证\n\n#### 业务领域动态模型（简历投递与邀约状态流转时序图）\n![业务涉及动态模型]({images_paths.get('sequence_diagram', '')})\n\n")

with open('README.md', 'w', encoding='utf-8') as f:
    f.write(readme_text)
print("README.md updated.")

print("Generating Word Document...")
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading1 = doc.styles['Heading 1']
heading1.font.name = 'Microsoft YaHei'
heading1.font.size = Pt(16)
heading1.font.color.rgb = RGBColor(0, 0, 0)
heading1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading2 = doc.styles['Heading 2']
heading2.font.name = 'Microsoft YaHei'
heading2.font.size = Pt(14)
heading2.font.color.rgb = RGBColor(0, 0, 0)
heading2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading3 = doc.styles['Heading 3']
heading3.font.name = 'Microsoft YaHei'
heading3.font.size = Pt(12)
heading3.font.color.rgb = RGBColor(0, 0, 0)
heading3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading4 = doc.styles['Heading 4']
heading4.font.name = 'Microsoft YaHei'
heading4.font.size = Pt(12)
heading4.font.color.rgb = RGBColor(0, 0, 0)
heading4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

lines = readme_text.split('\\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    
    if line.startswith('![') and '](' in line:
        img_path = line.split('](')[1].split(')')[0]
        if os.path.exists(img_path):
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            rt = doc.paragraphs[-1].add_run()
            try:
                rt.add_picture(img_path, width=Inches(5.5))
            except Exception as e:
                print(f"Could not add picture: {e}")
            
            lbl = line.split('![')[1].split(']')[0]
            p = doc.add_paragraph(lbl)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.style.font.size = Pt(10)
            except:
                pass
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        try:
            for run in p.runs:
                run.font.italic = True
        except:
            pass
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif len(line) > 0 and line[0].isdigit() and '. ' in line[:4]:
        idx = line.find('. ')
        doc.add_paragraph(line[idx+2:], style='List Number')
    elif line == '---' or line.startswith('```'):
        pass
    else:
        clean_line = line.replace('**', '').replace('`', '')
        doc.add_paragraph(clean_line)
        
    i += 1

doc_path = '职引未来-系统设计与评估报告.docx'
doc.save(doc_path)
print(f"Successfully generated {doc_path}")

