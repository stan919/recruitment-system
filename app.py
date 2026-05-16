from functools import lru_cache
"""
职引未来 - 高校毕业生就业服务平台
"""
import load_env

from flask import Flask, render_template, jsonify, request, session, redirect, url_for, send_file, Response
from flask_compress import Compress
from flask_cors import CORS
from markupsafe import escape
from datetime import timedelta, datetime
from collections import defaultdict, Counter
import config
import json
import io
import math
import re
import requests
import zipfile
import xml.etree.ElementTree as ET
import PyPDF2
from contextlib import contextmanager
from werkzeug.utils import secure_filename
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:
    Document = None

# 导入用户数据库模型
from models_user import User, get_session as UserSession, Base as UserBase
from models_job import JobPosition
from models_resume import Resume, JobApplication, get_session as ResumeSession, init_db as init_resume_db
from models_company import Company, AuditLog, init_db as init_company_db
from auth import AuthManager

# 构建数据库 URL
DATABASE_URL = f"mysql+pymysql://{config.DATABASE_CONFIG['user']}:{config.DATABASE_CONFIG['password']}@{config.DATABASE_CONFIG['host']}:{config.DATABASE_CONFIG['port']}/{config.DATABASE_CONFIG['database']}?charset=utf8mb4"

# 创建数据库引擎（单例模式）
from sqlalchemy import create_engine, and_, inspect, text, func
from sqlalchemy.orm import sessionmaker, scoped_session

engine = create_engine(
    DATABASE_URL,
    pool_size=20,
    max_overflow=10,
    pool_recycle=3600,
    pool_pre_ping=True,  # 连接前检测有效性
    echo=False
)
SessionLocal = scoped_session(sessionmaker(bind=engine))

# 数据库连接池监控
def get_db_pool_status():
    """获取数据库连接池状态"""
    pool = engine.pool
    return {
        'pool_size': pool.size(),
        'checked_in': pool.checkedin(),
        'checked_out': pool.checkedout(),
        'overflow': pool.overflow(),
        'invalidated': getattr(pool, '_invalidated', 0)
    }

# 配置日志
import logging
from logging.handlers import RotatingFileHandler
import os

# 确保日志目录存在
log_dir = 'logs'
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

SENSITIVE_WORDS_FILE = os.path.join(log_dir, 'sensitive_words.json')

# 配置日志处理器
file_handler = RotatingFileHandler(
    os.path.join(log_dir, 'app.log'),
    maxBytes=10*1024*1024,  # 10MB
    backupCount=10,
    encoding='utf-8'
)
file_handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
))
file_handler.setLevel(logging.INFO)

# 添加控制台处理器
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))

app = Flask(__name__)
# 启用 Gzip 压缩，大幅度提升接口传输性能
Compress(app)
# 对后端静态资源(CSS,JS,图片)默认使用一年的强缓存
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = timedelta(days=365)
APP_VERSION = '1.0.3'
app.config['SECRET_KEY'] = config.FLASK_CONFIG['SECRET_KEY']
app.config['DEBUG'] = config.FLASK_CONFIG['DEBUG']
import logging
# 关闭 Werkzeug 原生的繁杂请求日志提升网络 I/O 性能
log = logging.getLogger('werkzeug')
log.setLevel(logging.ERROR)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制文件大小为 16MB
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)  # 缩短Session有效期
app.config['SESSION_COOKIE_HTTPONLY'] = True  # 防止XSS读取Cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF防护

# 添加日志配置
if app.config['DEBUG']:
    app.logger.addHandler(console_handler)
else:
    app.logger.addHandler(file_handler)
    app.logger.setLevel(logging.INFO)

# 启用 CORS
CORS(app)

# CSRF 保护已在 Flask 表单中通过 {{ form.csrf_token }} 实现
# API 使用 Session 认证，不需要额外的 CSRF 保护
# csrf = CSRFProtect()
# csrf.init_app(app)

# 初始化数据库表
try:
    UserBase.metadata.create_all(engine)
    init_resume_db()
    init_company_db()
    app.logger.info("数据库表初始化完成")
except Exception as e:
    app.logger.warning("数据库初始化警告：%s", e)


def ensure_database_compatibility():
    """兼容历史库结构：补齐代码依赖但旧库中缺失的关键字段。"""
    try:
        db_inspector = inspect(engine)
        columns_map = {
            table: {col['name'] for col in db_inspector.get_columns(table)}
            for table in ['companies', 'conversations', 'messages', 'job_applications']
            if db_inspector.has_table(table)
        }

        alter_sql = []

        company_columns = columns_map.get('companies', set())
        if 'recruitment_manifesto' not in company_columns:
            alter_sql.append(
                "ALTER TABLE companies ADD COLUMN recruitment_manifesto TEXT NULL COMMENT '招聘宣言'"
            )

        conversation_columns = columns_map.get('conversations', set())
        if 'company_name' not in conversation_columns:
            alter_sql.append(
                "ALTER TABLE conversations ADD COLUMN company_name VARCHAR(200) NOT NULL DEFAULT '' COMMENT '公司名称'"
            )
        if 'job_name' not in conversation_columns:
            alter_sql.append(
                "ALTER TABLE conversations ADD COLUMN job_name VARCHAR(100) NULL COMMENT '职位名称'"
            )

        message_columns = columns_map.get('messages', set())
        if 'message_type' not in message_columns:
            alter_sql.append(
                "ALTER TABLE messages ADD COLUMN message_type VARCHAR(20) NULL COMMENT '消息类型'"
            )

        application_columns = columns_map.get('job_applications', set())
        if 'applicant_name' not in application_columns:
            alter_sql.append(
                "ALTER TABLE job_applications ADD COLUMN applicant_name VARCHAR(100) NULL COMMENT '申请时姓名快照'"
            )
        if 'applicant_phone' not in application_columns:
            alter_sql.append(
                "ALTER TABLE job_applications ADD COLUMN applicant_phone VARCHAR(30) NULL COMMENT '申请时手机号快照'"
            )
        if 'applicant_email' not in application_columns:
            alter_sql.append(
                "ALTER TABLE job_applications ADD COLUMN applicant_email VARCHAR(120) NULL COMMENT '申请时邮箱快照'"
            )

        if alter_sql:
            with engine.begin() as conn:
                for stmt in alter_sql:
                    conn.execute(text(stmt))
            app.logger.info("数据库兼容性修复完成，共执行 %s 条 DDL", len(alter_sql))

        # 旧数据回填：将历史申请中为空的快照字段补齐并固定，避免后续受简历改动影响
        if db_inspector.has_table('job_applications') and db_inspector.has_table('resumes'):
            with engine.begin() as conn:
                conn.execute(text("""
                    UPDATE job_applications ja
                    LEFT JOIN resumes r ON ja.resume_id = r.id
                    LEFT JOIN users u ON ja.user_id = u.id
                    SET
                        ja.applicant_name = COALESCE(ja.applicant_name, r.name, u.username),
                        ja.applicant_phone = COALESCE(ja.applicant_phone, r.phone, u.phone),
                        ja.applicant_email = COALESCE(ja.applicant_email, r.email, u.email)
                    WHERE
                        ja.applicant_name IS NULL
                        OR ja.applicant_phone IS NULL
                        OR ja.applicant_email IS NULL
                """))
    except Exception as e:
        app.logger.warning("数据库兼容性修复警告：%s", e)


ensure_database_compatibility()


def load_sensitive_words():
    """从本地文件加载敏感词列表。"""
    if not os.path.exists(SENSITIVE_WORDS_FILE):
        return []

    try:
        with open(SENSITIVE_WORDS_FILE, 'r', encoding='utf-8') as f:
            words = json.load(f)
        if not isinstance(words, list):
            return []
        return words
    except Exception:
        return []


def save_sensitive_words(words):
    """将敏感词列表保存到本地文件。"""
    try:
        with open(SENSITIVE_WORDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(words, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


# 速率限制 - 记录登录尝试
login_attempts = defaultdict(list)
api_rate_limits = defaultdict(list)  # API速率限制
RATE_LIMIT_WINDOW = 300  # 5 分钟
MAX_LOGIN_ATTEMPTS = 10  # 最多 10 次尝试
MAX_TRACKED_LOGIN_KEYS = 5000  # 内存中最多保留的限流键数量
MAX_CHAT_MESSAGE_LENGTH = 2000
MAX_COMPANY_NAME_LENGTH = 200
MAX_JOB_NAME_LENGTH = 100
ALLOWED_CHAT_MESSAGE_TYPES = {'text', 'voice', 'emoji'}
INTERVIEW_CONFIRMATION_MARKER = '[INTERVIEW_CONFIRMED_BY_USER]'

# API速率限制配置
API_RATE_LIMITS = {
    '/api/resume/upload': {'max_requests': 10, 'window': 3600},  # 每小时10次上传
    '/api/account/avatar': {'max_requests': 20, 'window': 3600},  # 每小时20次头像上传
    '/api/chat/send': {'max_requests': 60, 'window': 60},  # 每分钟60条消息
    '/api/application/submit': {'max_requests': 20, 'window': 3600},  # 每小时20次申请
}

# 管理控制台统计缓存，降低重复 count 查询带来的首屏延迟
ADMIN_DASHBOARD_CACHE = {
    'expires_at': None,
    'payload': None
}
ADMIN_DASHBOARD_CACHE_TTL_SECONDS = 20

# 洞察数据缓存（减少重计算）
INSIGHTS_CACHE = {
    'profession_data': {},  # {profession_key: {'expires_at': ..., 'data': ...}}
    'overall_skill_cloud': {'expires_at': None, 'data': None},
    'overall_data': {'expires_at': None, 'data': None}
}
INSIGHTS_CACHE_TTL_SECONDS = 300  # 5分钟缓存


def invalidate_admin_dashboard_cache():
    """管理员控制台统计缓存失效。"""
    ADMIN_DASHBOARD_CACHE['expires_at'] = None
    ADMIN_DASHBOARD_CACHE['payload'] = None


# ==================== 权限检查装饰器 ====================

def require_login(f):
    """需要登录的装饰器"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请先登录'}), 401
        return f(*args, **kwargs)
    return decorated_function


def require_user_role(f):
    """仅普通用户可访问的装饰器"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请先登录'}), 401
        
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if user and (user.is_company_admin or user.is_admin):
                return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
        return f(*args, **kwargs)
    return decorated_function


# ==================== 全局异常处理器 ====================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'success': False, 'message': '资源不存在'}), 404


@app.errorhandler(500)
def internal_error(error):
    app.logger.exception("服务器内部错误")
    return jsonify({'success': False, 'message': '服务器内部错误，请稍后重试'}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'success': False, 'message': '文件过大'}), 413


# ==================== 请求日志中间件 ====================

@app.before_request
def log_request_info():
    """记录请求信息用于监控和速率限制"""
    if request.path.startswith('/api/'):
        request.start_time = datetime.now()
        
        # API速率限制检查
        client_ip = get_client_ip()
        rate_key = f"{client_ip}:{request.path}"
        
        # 检查是否有针对该路径的速率限制
        for path_prefix, limit_config in API_RATE_LIMITS.items():
            if request.path.startswith(path_prefix):
                now = datetime.now()
                window = limit_config['window']
                max_requests = limit_config['max_requests']
                
                # 清理过期记录
                api_rate_limits[rate_key] = [
                    ts for ts in api_rate_limits[rate_key]
                    if (now - ts).total_seconds() < window
                ]
                
                # 检查是否超过限制
                if len(api_rate_limits[rate_key]) >= max_requests:
                    app.logger.warning(f"API速率限制: {rate_key}")
                    return jsonify({
                        'success': False,
                        'message': '请求过于频繁，请稍后再试'
                    }), 429
                
                # 记录本次请求
                api_rate_limits[rate_key].append(now)
                break


@app.after_request
def log_response_info(response):
    """记录响应时间和状态码"""
    if hasattr(request, 'start_time'):
        duration = (datetime.now() - request.start_time).total_seconds() * 1000
        app.logger.info(
            f"{request.method} {request.path} - {response.status_code} - {duration:.0f}ms"
        )
    return response


# ==================== 工具函数 ====================

@contextmanager
def get_db_session():
    """数据库会话上下文管理器"""
    db_session = SessionLocal()
    try:
        yield db_session
    except Exception:
        db_session.rollback()
        raise
    finally:
        db_session.close()


def get_actual_city(city_code):
    """获取实际城市名"""
    return config.CITY_MAPPING.get(city_code, city_code)


def get_profession_keyword(profession):
    """获取专业关键词"""
    return config.PROFESSION_KEYWORDS.get(profession, '')


def get_redirect_by_session_role():
    """根据当前登录模式返回默认跳转地址。"""
    if session.get('is_admin'):
        return '/admin'
    if session.get('is_company'):
        return '/company'
    return '/'


def coerce_bool(value):
    """将请求参数中的布尔值安全归一化。"""
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    if isinstance(value, str):
        return value.strip().lower() in {'1', 'true', 'yes', 'on'}
    return False


def normalize_company_name(value):
    """统一公司名称比较口径，避免空白/大小写差异导致误判。"""
    return (value or '').strip().casefold()


def find_company_record_by_name(db_session, company_name):
    """按公司名查找公司记录（先精确，再标准化匹配）。"""
    raw_name = (company_name or '').strip()
    if not raw_name:
        return None

    company = db_session.query(Company).filter_by(company_name=raw_name).first()
    if company:
        return company

    normalized = normalize_company_name(raw_name)
    if not normalized:
        return None

    return db_session.query(Company).filter(
        func.lower(func.trim(Company.company_name)) == normalized
    ).order_by(Company.id.asc()).first()


def get_client_ip():
    """尽可能获取真实客户端 IP。"""
    forwarded_for = (request.headers.get('X-Forwarded-For') or '').strip()
    if forwarded_for:
        return forwarded_for.split(',')[0].strip()

    real_ip = (request.headers.get('X-Real-IP') or '').strip()
    return real_ip or (request.remote_addr or 'unknown')


def build_login_attempt_key(username, client_ip):
    """构造登录限流键：用户名 + IP。"""
    return f"{username.lower()}|{client_ip}"


def cleanup_login_attempts(now):
    """清理过期限流记录并控制字典大小（优化版：批量处理）。"""
    expired_keys = []
    
    for key, timestamps in login_attempts.items():
        # 过滤有效时间戳
        recent = [ts for ts in timestamps if (now - ts).total_seconds() < RATE_LIMIT_WINDOW]
        
        if recent:
            login_attempts[key] = recent
        else:
            expired_keys.append(key)
    
    # 批量删除过期键
    for key in expired_keys:
        login_attempts.pop(key, None)
    
    # 控制字典大小
    if len(login_attempts) > MAX_TRACKED_LOGIN_KEYS:
        overflow = len(login_attempts) - MAX_TRACKED_LOGIN_KEYS
        oldest_keys = sorted(
            login_attempts.keys(),
            key=lambda k: login_attempts[k][-1] if login_attempts[k] else now
        )[:overflow]
        for key in oldest_keys:
            login_attempts.pop(key, None)


def extract_docx_preview_html(file_path):
    """提取 docx 文本并转换为轻量 HTML 预览。"""
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    with zipfile.ZipFile(file_path, 'r') as archive:
        xml_content = archive.read('word/document.xml')

    root = ET.fromstring(xml_content)
    paragraphs = []

    for paragraph in root.findall('.//w:p', namespace):
        chunks = []
        for node in paragraph.findall('.//w:t', namespace):
            if node.text:
                chunks.append(node.text)

        line = ''.join(chunks).strip()
        if line:
            paragraphs.append(f'<p>{escape(line)}</p>')

    if not paragraphs:
        return '<p>文档内容为空</p>'

    return ''.join(paragraphs)


def extract_doc_preview_html(file_path):
    """使用 mammoth 将 doc/docx 转换为 HTML。"""
    import mammoth

    with open(file_path, 'rb') as word_file:
        result = mammoth.convert_to_html(word_file)

    return (result.value or '').strip() or '<p>文档内容为空</p>'


def extract_text_from_pdf(file_path):
    """从 PDF 中提取文本（尽量容错）。"""
    try:
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    txt = page.extract_text() or ''
                except Exception:
                    txt = ''
                if txt:
                    text_parts.append(txt)
        return '\n'.join(text_parts).strip()
    except Exception:
        app.logger.exception('PDF 文本提取失败')
        return ''


def extract_text_from_docx(file_path):
    """从 docx/doc 转为 HTML 再提取纯文本，作为后端分析的文本来源。"""
    try:
        html = extract_doc_preview_html(file_path)
        # 简单去除 HTML 标签
        txt = re.sub(r'<[^>]+>', '\n', html)
        txt = re.sub(r'\n{2,}', '\n', txt)
        return txt.strip()
    except Exception:
        app.logger.exception('DOCX 文本提取失败')
        return ''


def extract_text_from_image(file_path):
    """从图片中提取文本；优先使用可选 OCR，失败则返回空字符串。"""
    try:
        from PIL import Image
    except Exception:
        app.logger.warning('图片解析依赖不可用，跳过 OCR')
        return ''

    try:
        image = Image.open(file_path)
        image.load()
    except Exception:
        app.logger.exception('图片文件打开失败')
        return ''

    try:
        import pytesseract
    except Exception:
        app.logger.warning('pytesseract 未安装，图片仅保留文件信息参与分析')
        return ''

    try:
        # 尽量提高中文/英文混排文本的识别率
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        return (text or '').strip()
    except Exception:
        app.logger.exception('图片 OCR 失败')
        return ''


def extract_text_from_file(file_path, ext):
    """根据扩展名提取文件文本，返回字符串或空字符串。"""
    ext = (ext or '').lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    if ext in ('docx', 'doc'):
        return extract_text_from_docx(file_path)
    if ext in ('png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif'):
        return extract_text_from_image(file_path)
    return ''


def extract_ai_reply(result):
    """兼容不同供应商的 AI 响应结构。"""
    reply = ''
    if isinstance(result, dict):
        choices = result.get('choices') or []
        if choices:
            first_choice = choices[0] if isinstance(choices[0], dict) else {}
            message = first_choice.get('message') or {}
            if isinstance(message, dict):
                reply = str(message.get('content', '')).strip()
            if not reply:
                text = first_choice.get('text')
                if text is not None:
                    reply = str(text).strip()

        if not reply:
            output = result.get('output') or {}
            if isinstance(output, dict):
                reply = str(output.get('text') or output.get('content') or '').strip()

        if not reply:
            reply = str(result.get('reply') or result.get('message') or '').strip()

    return reply


def extract_provider_error(response):
    """从提供方响应中提取可读错误信息。"""
    if response is None:
        return None

    try:
        body = response.json()
    except ValueError:
        return (response.text or '').strip() or None

    if isinstance(body, dict):
        error = body.get('error')
        if isinstance(error, dict):
            return error.get('message') or error.get('detail') or str(error)
        return body.get('message') or body.get('detail') or str(body)

    return str(body)


def build_ai_request_candidates(base_url, model, user_msg, ai_config, system_prompt=None):
    """构建兼容 DashScope compatible-mode 的单一路径请求。"""
    if not system_prompt:
        system_prompt = (
            '你是一个专业的招聘与职业规划 AI 助手。'
            '回答要具体、可执行、简洁，优先给出简历优化、岗位匹配、面试准备和职业建议。'
        )
    temperature = ai_config.get('TEMPERATURE', 0.7)
    max_tokens = ai_config.get('MAX_TOKENS', 1200)

    return [
        {
            'url': f'{base_url}/chat/completions',
            'payload': {
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_msg},
                ],
                'temperature': temperature,
                'max_tokens': max_tokens,
            },
        },
    ]


def parse_ai_json_reply(text):
    """从模型回复中提取 JSON。"""
    raw_text = str(text or '').strip()
    if not raw_text:
        return None

    if raw_text.startswith('```'):
        raw_text = re.sub(r'^```(?:json)?\s*', '', raw_text, flags=re.IGNORECASE).strip()
        raw_text = re.sub(r'\s*```$', '', raw_text).strip()

    candidates = [raw_text]
    match = re.search(r'\{[\s\S]*\}', raw_text)
    if match:
        candidates.insert(0, match.group(0))

    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            continue
    return None


def summarize_applications_for_ai(applications):
    """把投递记录整理成可用于 AI 分析的摘要。"""
    applications = list(applications or [])
    status_counter = Counter()
    recent_jobs = []

    for application in applications:
        status = str(getattr(application, 'status', '') or '未知').strip() or '未知'
        status_counter[status] += 1
        job_name = str(getattr(application, 'job_name', '') or '').strip()
        company_name = str(getattr(application, 'company_name', '') or '').strip()
        if job_name or company_name:
            recent_jobs.append(' / '.join([item for item in [job_name, company_name] if item]))

    return {
        'total': len(applications),
        'status_counts': dict(status_counter),
        'recent_jobs': recent_jobs[:5],
    }


def build_resume_ai_fallback_analysis(source_text, file_name, template_key, resume_data=None):
    """当后端 AI 不可用时，生成一个结构化兜底分析。"""
    profile_map = {
        'general': {
            'base_score': 68,
            'headline': '先补结构，再提亮点',
            'summary': '先确认简历是否容易被快速扫描，再补项目成果、技能关键词和量化证据。',
            'rewrite': '把“负责项目开发”改成“主导某模块开发，接入 3 个关键功能，帮助团队将响应效率提升 30%”。',
        },
        'data': {
            'base_score': 72,
            'headline': '数据岗先看可量化分析能力',
            'summary': '数据岗更看重指标表达、分析方法和工具栈，建议把 SQL / Excel / Python / BI 相关经历写实。',
            'rewrite': '把“做过数据分析”改成“基于 SQL 和 Python 分析用户行为，输出 3 份周报并定位了 2 个关键流失节点”。',
        },
        'product': {
            'base_score': 70,
            'headline': '产品岗要先讲清目标与结果',
            'summary': '产品岗更看重需求拆解、协作推进和结果验证，建议把用户价值、迭代节奏和 A/B 结果补完整。',
            'rewrite': '把“参与产品设计”改成“独立推动某功能从需求评审到上线，协同设计与研发完成 2 轮迭代，转化率提升 12%”。',
        },
        'engineering': {
            'base_score': 74,
            'headline': '研发岗优先突出技术栈和交付能力',
            'summary': '研发岗建议明确语言、框架、架构、性能优化和协作方式，把“做过”改成“交付了什么”。',
            'rewrite': '把“负责后端开发”改成“基于 Flask 完成接口重构，优化查询逻辑并将接口响应时间降低 28%”。',
        },
    }

    profile = profile_map.get(template_key, profile_map['general'])
    text = str(source_text or '').strip()
    resume_data = resume_data or {}
    lower_text = f"{text} {file_name or ''} {resume_data.get('work_experience', '')} {resume_data.get('project_experience', '')} {resume_data.get('skills', '')}".lower()

    score = profile['base_score']
    strengths = []
    issues = []

    if text:
        if len(text) > 280:
            score += 12
            strengths.append('补充内容较完整，适合进一步精修')
        elif len(text) > 120:
            score += 7
            strengths.append('已有较完整的文字材料，可以继续强化成果')
        else:
            score -= 6
            issues.append('补充内容偏少，建议增加项目和经历细节')
    else:
        issues.append('当前补充文本较少，建议把简历重点贴到工作台里')

    if resume_data.get('name'):
        score += 3
        strengths.append('基础个人信息已包含在简历中')
    if resume_data.get('education'):
        score += 3
        strengths.append('教育背景信息完整')
    if resume_data.get('work_experience') or resume_data.get('project_experience'):
        score += 5
        strengths.append('有工作或项目经历可供继续提炼')
    else:
        issues.append('建议补充项目经历或实习经历，增强可读性')

    if re.search(r'\d+(?:\.\d+)?%|\d+倍|\d+万|\d+人|\d+个|\d+小时', text or ''):
        score += 10
        strengths.append('出现了量化表达，比较利于面试官判断成果')
    else:
        issues.append('缺少量化结果，建议补上百分比、人数或周期')

    keyword_rules = {
        'general': [r'项目', r'实习', r'经历', r'负责', r'主导', r'成果', r'技能', r'协作'],
        'data': [r'sql', r'python', r'pandas', r'excel', r'bi', r'tableau', r'power bi', r'sklearn'],
        'product': [r'产品', r'用户', r'需求', r'a/b', r'原型', r'迭代', r'转化', r'增长'],
        'engineering': [r'java', r'python', r'js', r'ts', r'react', r'vue', r'flask', r'django', r'docker', r'api', r'mysql', r'redis', r'linux'],
    }
    keyword_hits = sum(1 for rule in keyword_rules['general'] if re.search(rule, lower_text))
    keyword_hits += sum(1 for rule in keyword_rules.get(template_key, []) if re.search(rule, lower_text))

    if keyword_hits >= 6:
        score += 8
        strengths.append('岗位关键词覆盖比较到位')
    elif keyword_hits >= 3:
        score += 3
        issues.append('关键词数量还可以再补一轮')
    else:
        issues.append('关键词密度偏低，建议围绕目标岗位补词')

    score = max(30, min(98, int(score)))
    priority = '高' if len(issues) >= 3 else '中' if len(issues) >= 2 else '低'

    if not strengths:
        strengths.append('内容可继续细化，适合后续迭代')
    if not issues:
        issues.append('当前材料已经有一定基础，可以开始做针对性优化')

    return {
        'score': score,
        'headline': profile['headline'],
        'summary': profile['summary'],
        'structure': '完整' if len(text) > 180 or resume_data.get('work_experience') else '待补充',
        'quant': '较强' if re.search(r'\d+(?:\.\d+)?%|\d+倍|\d+万|\d+人|\d+个|\d+小时', text or '') else '待加强',
        'keywords': '较高' if keyword_hits >= 4 else '待加强',
        'priority': priority,
        'strengths': strengths[:4],
        'issues': issues[:4],
        'rewrite': profile['rewrite'],
    }


def normalize_resume_ai_analysis(parsed_analysis, fallback_analysis):
    """把模型输出归一到前端需要的字段。"""
    analysis = dict(fallback_analysis or {})
    if not isinstance(parsed_analysis, dict):
        return analysis

    for key in ['headline', 'summary', 'structure', 'quant', 'keywords', 'priority', 'rewrite']:
        value = parsed_analysis.get(key)
        if isinstance(value, str) and value.strip():
            analysis[key] = value.strip()

    score = parsed_analysis.get('score')
    if isinstance(score, (int, float)):
        analysis['score'] = max(0, min(100, int(score)))

    strengths = parsed_analysis.get('strengths')
    if isinstance(strengths, list) and strengths:
        analysis['strengths'] = [str(item).strip() for item in strengths if str(item).strip()][:5]

    issues = parsed_analysis.get('issues')
    if isinstance(issues, list) and issues:
        analysis['issues'] = [str(item).strip() for item in issues if str(item).strip()][:5]

    return analysis


def build_resume_ai_docx(analysis, file_name, template_key):
    """把简历分析结果打包为可下载的 Word 文档。"""
    if Document is None:
        raise RuntimeError('docx_not_available')

    def safe_text(value, default=''):
        text = str(value or '').strip()
        return text or default

    def safe_list(value):
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        if isinstance(value, str) and value.strip():
            return [value.strip()]
        return []

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    except Exception:
        pass

    title = doc.add_heading('AI 简历优化建议', level=1)
    title.alignment = 1

    meta_line = f"文件：{safe_text(file_name, '未命名简历')} | 视角：{safe_text(template_key, '通用')} | 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    doc.add_paragraph(meta_line)

    headline = safe_text(analysis.get('headline'), '简历分析概要')
    summary = safe_text(analysis.get('summary'), '暂无摘要')
    doc.add_heading('摘要', level=2)
    doc.add_paragraph(f"{headline}。{summary}")

    doc.add_heading('结构化诊断', level=2)
    structure = safe_text(analysis.get('structure'), '待补充')
    quant = safe_text(analysis.get('quant'), '待加强')
    keywords = safe_text(analysis.get('keywords'), '待加强')
    priority = safe_text(analysis.get('priority'), '中')
    doc.add_paragraph(f"结构完整度：{structure}")
    doc.add_paragraph(f"量化表达：{quant}")
    doc.add_paragraph(f"关键词贴合：{keywords}")
    doc.add_paragraph(f"优化优先级：{priority}")

    strengths = safe_list(analysis.get('strengths'))
    issues = safe_list(analysis.get('issues'))

    doc.add_heading('亮点', level=2)
    if strengths:
        for item in strengths:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph('暂无明显亮点，可从项目结果和量化数据补充。')

    doc.add_heading('待补强点', level=2)
    if issues:
        for item in issues:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph('暂无明显短板，可继续优化表达和关键词布局。')

    rewrite = safe_text(analysis.get('rewrite'), '')
    if rewrite:
        doc.add_heading('改写示例', level=2)
        doc.add_paragraph(rewrite)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def looks_like_realtime_model(model_name):
    """识别明显需要实时通道而不是普通 HTTP chat/completions 的模型。"""
    normalized = str(model_name or '').strip().lower()
    return bool(normalized) and ('realtime' in normalized or 'live' in normalized)


@app.before_request
def enforce_admin_session_mode():
    """限制管理后台仅允许“超级管理员登录态”访问。"""
    path = request.path or ''
    if path == '/admin' or path.startswith('/api/admin/'):
        if 'user_id' not in session:
            if path == '/admin':
                return redirect('/login')
            return jsonify({'success': False, 'message': '未登录'}), 401

        if not session.get('is_admin', False):
            if path == '/admin':
                return redirect('/login?message=请使用超级管理员方式登录')
            return jsonify({'success': False, 'message': '请使用超级管理员方式登录'}), 403


# ==================== 认证相关 API ====================

@app.route('/health')
def health_check():
    """健康检查 API - 增强版"""
    try:
        # 检查数据库连接
        with engine.connect() as conn:
            conn.execute(text('SELECT 1'))
        db_status = 'connected'
    except Exception as e:
        app.logger.error(f"数据库健康检查失败: {e}")
        db_status = 'disconnected'
    
    # 获取连接池状态
    pool_status = get_db_pool_status()
    
    # 计算缓存统计
    cache_stats = {
        'profession_data_entries': len(INSIGHTS_CACHE['profession_data']),
        'overall_skill_cloud_cached': INSIGHTS_CACHE['overall_skill_cloud']['expires_at'] is not None,
        'overall_data_cached': INSIGHTS_CACHE['overall_data']['expires_at'] is not None
    }
    
    status = 'healthy' if db_status == 'connected' else 'degraded'
    
    return jsonify({
        'status': status,
        'timestamp': datetime.now().isoformat(),
        'version': APP_VERSION,
        'database': {
            'status': db_status,
            'pool': pool_status
        },
        'cache': cache_stats,
        'rate_limits': {
            'login_attempts_keys': len(login_attempts),
            'api_rate_limit_keys': len(api_rate_limits)
        }
    }), 200 if status == 'healthy' else 503

@app.route('/login')
def login():
    """登录页面"""
    return render_template('login.html')


@app.route('/register')
def register():
    """注册页面"""
    return render_template('register.html')


@app.route('/api/login', methods=['POST'])
def api_login():
    """用户登录 API"""
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    password = data.get('password') or ''
    is_company = coerce_bool(data.get('is_company', False))
    is_admin = coerce_bool(data.get('is_admin', False))
    remember = coerce_bool(data.get('remember', False))

    if is_company and is_admin:
        return jsonify({'success': False, 'message': '请只选择一种登录方式'}), 400
    
    if not username or not password:
        return jsonify({'success': False, 'message': '请输入用户名和密码'}), 400
    
    # 速率限制检查
    now = datetime.now()
    cleanup_login_attempts(now)
    client_ip = get_client_ip()
    rate_key = build_login_attempt_key(username, client_ip)
    
    # 检查是否超过限制
    if len(login_attempts.get(rate_key, [])) >= MAX_LOGIN_ATTEMPTS:
        app.logger.warning(f"登录失败次数过多: {username} from {client_ip}")
        return jsonify({'success': False, 'message': '尝试次数过多，请稍后再试'}), 429
    
    user, error = AuthManager.authenticate(username, password)
    
    if user:
        # 检查用户是否被禁用
        if not user.is_active:
            app.logger.warning(f"尝试登录已被禁用的账户: {username}")
            return jsonify({'success': False, 'message': '账户已被禁用'}), 401

        # 登录成功后清空该账号的失败计数，避免误触发限流
        login_attempts.pop(rate_key, None)

        # 如果是企业登录，检查是否为企业管理员
        if is_company and not user.is_company_admin:
            app.logger.warning(f"非企业管理员尝试企业登录: {username}")
            return jsonify({'success': False, 'message': '该账号不是企业管理员'}), 401
        
        # 如果是管理员登录，检查是否为超级管理员
        if is_admin and not user.is_admin:
            app.logger.warning(f"非超级管理员尝试管理员登录: {username}")
            return jsonify({'success': False, 'message': '该账号不是超级管理员'}), 401

        # 超级管理员账号必须通过"超级管理员登录"入口进入后台
        if user.is_admin and not is_admin:
            app.logger.warning(f"超级管理员未使用管理员入口登录: {username}")
            return jsonify({'success': False, 'message': '请使用超级管理员方式登录'}), 401

        # 企业管理员账号必须通过"企业登录"入口进入后台
        if user.is_company_admin and not is_company:
            app.logger.warning(f"企业管理员未使用企业入口登录: {username}")
            return jsonify({'success': False, 'message': '请使用企业登录方式登录'}), 401
        
        session['user_id'] = user.id
        session['username'] = user.username
        session['email'] = user.email
        session['is_company'] = is_company
        session['is_admin'] = is_admin
        session.permanent = remember
        
        # 记录成功登录审计日志
        app.logger.info(
            f"用户登录成功: {username} (ID:{user.id}) from {client_ip} "
            f"type={'admin' if is_admin else 'company' if is_company else 'user'}"
        )
        
        # 根据登录类型重定向到不同页面
        if is_admin:
            redirect_url = '/admin'
        elif is_company:
            redirect_url = '/company'
        else:
            redirect_url = '/'
        
        return jsonify({
            'success': True,
            'message': '登录成功',
            'user': user.to_dict(),
            'redirect': redirect_url
        })
    else:
        # 仅对失败登录计数
        login_attempts[rate_key].append(now)
        app.logger.warning(f"登录失败: {username} from {client_ip} - {error}")
        return jsonify({'success': False, 'message': error or '登录失败'}), 401


@app.route('/api/register', methods=['POST'])
def api_register():
    """用户注册 API"""
    data = request.get_json(silent=True) or {}
    username = data.get('username', '').strip()
    email = data.get('email', '').strip()
    phone = data.get('phone', '').strip()
    password = data.get('password', '')
    
    if not username or not email or not password:
        return jsonify({'success': False, 'message': '请填写必填项'}), 400
    
    if len(username) < 3 or len(username) > 20:
        return jsonify({'success': False, 'message': '用户名长度必须在 3-20 个字符之间'}), 400
    
    if len(password) < 8:
        return jsonify({'success': False, 'message': '密码长度至少 8 位'}), 400
    
    user, error = AuthManager.create_user(username, email, password, phone)
    
    if user:
        app.logger.info(f"新用户注册: {username} from {get_client_ip()}")
        return jsonify({'success': True, 'message': '注册成功'})
    else:
        app.logger.warning(f"注册失败: {username} - {error}")
        return jsonify({'success': False, 'message': error or '注册失败'}), 400


@app.route('/api/logout', methods=['POST'])
def api_logout():
    """用户登出 API"""
    session.clear()
    return jsonify({'success': True, 'message': '已退出登录'})


@app.route('/api/user')
def api_user():
    """获取当前用户信息"""
    if 'user_id' in session:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if user:
                return jsonify({
                    'logged_in': True,
                    'user': user.to_dict(),
                    'session_flags': {
                        'is_company': bool(session.get('is_company', False)),
                        'is_admin': bool(session.get('is_admin', False))
                    }
                })
    
    return jsonify({'logged_in': False}), 401


# ==================== 主页 ====================

@app.route('/')
def index():
    """首页 - 展示宣传主页"""
    if session.get('is_admin'):
        return redirect('/admin')

    if 'user_id' in session:
        with get_db_session() as db_session:
            current_user = db_session.query(User).filter_by(id=session['user_id']).first()
            if current_user and current_user.is_company_admin:
                return redirect('/company')

    # 获取热门职位（用于搜索）
    jobs = []
    try:
        with get_db_session() as db_session:
            jobs = db_session.query(JobPosition).filter_by(status='active').limit(20).all()
    except Exception as e:
        app.logger.warning("获取首页职位失败：%s", e)
    
    return render_template('home.html', jobs=jobs)


@app.route('/dashboard')
def dashboard():
    """主页 - 跳转到宣传主页"""
    return redirect(url_for('index'))


@app.route('/insights')
def insights():
    """数据洞察页面"""
    return render_template('insights.html')


@app.route('/api/insights/professions')
def insights_professions():
    """获取可筛选的专业列表。"""
    options = [
        {'key': key, 'name': value}
        for key, value in config.PROFESSION_KEYWORDS.items()
    ]
    return jsonify({'success': True, 'professions': options})


@app.route('/api/insights/profession-data')
def insights_profession_data():
    """按专业返回数据洞察图表所需数据。"""
    profession = request.args.get('profession', '').strip()
    keyword_text = get_profession_keyword(profession)

    if not profession or not keyword_text:
        return jsonify({'success': False, 'message': '专业参数无效'}), 400

    # 检查缓存
    now = datetime.now()
    cache_entry = INSIGHTS_CACHE['profession_data'].get(profession)
    if cache_entry and cache_entry['expires_at'] and now < cache_entry['expires_at']:
        app.logger.info(f"命中专业洞察缓存: {profession}")
        return jsonify({'success': True, **cache_entry['data'], 'cached': True})

    def percentile(values, p):
        if not values: return 0.0
        import statistics
        return float(statistics.quantiles(values, n=100, method='inclusive')[int(p * 100) - 1]) if 0 < p < 1 else float(max(values) if p == 1 else min(values))

    def avg_salary(min_salary, max_salary):
        nums = [
            float(v)
            for v in [min_salary, max_salary]
            if isinstance(v, (int, float)) and v > 0
        ]
        if not nums:
            return None
        return sum(nums) / len(nums)

    # 技能过滤词：显式排除“非技能类描述词”
    blocked_exact = {
        '薪资', '经验', '学历', '待遇', '要求', '应届生', '大专', '硕士', '本科', '以下',
        '运营', '人力', '市场', '专员', '分析师', '岗位', '职位', '工作', '相关', '优先', '不限'
    }
    blocked_contains = {
        '薪资', '经验', '学历', '待遇', '应届', '大专', '硕士', '本科', '以下',
        '专员', '分析师', '运营', '人力', '市场', '招聘', '岗位', '职位'
    }

    def build_mock_payload(profession_key):
        """当 ai / ba 无真实数据时，返回可视化模拟数据兜底。"""
        if profession_key == 'ai':
            city_distribution = [
                {'city': '北京', 'count': 720},
                {'city': '上海', 'count': 650},
                {'city': '深圳', 'count': 610},
                {'city': '杭州', 'count': 520},
                {'city': '广州', 'count': 430},
                {'city': '成都', 'count': 310}
            ]
            salary_by_city = [
                {'city': '北京', 'avg': 28300.0},
                {'city': '上海', 'avg': 27100.0},
                {'city': '深圳', 'avg': 26600.0},
                {'city': '杭州', 'avg': 25200.0},
                {'city': '广州', 'avg': 23600.0},
                {'city': '成都', 'avg': 21400.0}
            ]
            salary_quantiles = [
                {'city': '北京', 'p25': 22000.0, 'p50': 28000.0, 'p75': 34000.0},
                {'city': '上海', 'p25': 21000.0, 'p50': 27000.0, 'p75': 33000.0},
                {'city': '深圳', 'p25': 20500.0, 'p50': 26500.0, 'p75': 32200.0},
                {'city': '杭州', 'p25': 19000.0, 'p50': 25000.0, 'p75': 30500.0},
                {'city': '广州', 'p25': 17800.0, 'p50': 23500.0, 'p75': 29200.0},
                {'city': '成都', 'p25': 16000.0, 'p50': 21200.0, 'p75': 26800.0}
            ]
            education_distribution = [
                {'name': '本科', 'value': 1420},
                {'name': '硕士', 'value': 1090},
                {'name': '大专', 'value': 560},
                {'name': '博士', 'value': 170}
            ]
            experience_distribution = [
                {'name': '不限', 'value': 460},
                {'name': '应届生', 'value': 350},
                {'name': '1-3 年', 'value': 1120},
                {'name': '3-5 年', 'value': 910},
                {'name': '5-10 年', 'value': 400}
            ]
            skill_cloud = [
                {'name': 'Python', 'value': 980},
                {'name': '机器学习', 'value': 930},
                {'name': '深度学习', 'value': 860},
                {'name': 'PyTorch', 'value': 790},
                {'name': 'TensorFlow', 'value': 740},
                {'name': 'NLP', 'value': 670},
                {'name': '计算机视觉', 'value': 640},
                {'name': '大模型', 'value': 620},
                {'name': 'LLM', 'value': 590},
                {'name': '特征工程', 'value': 560},
                {'name': '模型部署', 'value': 520},
                {'name': '数据挖掘', 'value': 490},
                {'name': 'Linux', 'value': 450},
                {'name': 'SQL', 'value': 430},
                {'name': '算法', 'value': 400}
            ]
        elif profession_key == 'ba':
            city_distribution = [
                {'city': '上海', 'count': 680},
                {'city': '北京', 'count': 590},
                {'city': '深圳', 'count': 520},
                {'city': '广州', 'count': 470},
                {'city': '杭州', 'count': 430},
                {'city': '南京', 'count': 390}
            ]
            salary_by_city = [
                {'city': '上海', 'avg': 17800.0},
                {'city': '北京', 'avg': 17200.0},
                {'city': '深圳', 'avg': 16800.0},
                {'city': '广州', 'avg': 15400.0},
                {'city': '杭州', 'avg': 14900.0},
                {'city': '南京', 'avg': 13800.0}
            ]
            salary_quantiles = [
                {'city': '上海', 'p25': 13200.0, 'p50': 17600.0, 'p75': 22200.0},
                {'city': '北京', 'p25': 12800.0, 'p50': 17000.0, 'p75': 21500.0},
                {'city': '深圳', 'p25': 12400.0, 'p50': 16600.0, 'p75': 20900.0},
                {'city': '广州', 'p25': 11600.0, 'p50': 15100.0, 'p75': 19500.0},
                {'city': '杭州', 'p25': 11200.0, 'p50': 14600.0, 'p75': 18800.0},
                {'city': '南京', 'p25': 10400.0, 'p50': 13600.0, 'p75': 17500.0}
            ]
            education_distribution = [
                {'name': '本科', 'value': 1480},
                {'name': '大专', 'value': 920},
                {'name': '硕士', 'value': 440},
                {'name': '不限', 'value': 240}
            ]
            experience_distribution = [
                {'name': '不限', 'value': 560},
                {'name': '应届生', 'value': 420},
                {'name': '1-3 年', 'value': 1150},
                {'name': '3-5 年', 'value': 720},
                {'name': '5-10 年', 'value': 230}
            ]
            skill_cloud = [
                {'name': '数据分析', 'value': 920},
                {'name': '商业分析', 'value': 860},
                {'name': '财务报表', 'value': 800},
                {'name': 'PowerBI', 'value': 740},
                {'name': 'Excel', 'value': 700},
                {'name': '项目管理', 'value': 660},
                {'name': '战略规划', 'value': 610},
                {'name': '流程优化', 'value': 580},
                {'name': '成本控制', 'value': 540},
                {'name': '预算管理', 'value': 500},
                {'name': 'SQL', 'value': 460},
                {'name': '沟通协调', 'value': 430},
                {'name': '跨部门协作', 'value': 390},
                {'name': '运营管理', 'value': 350},
                {'name': '市场洞察', 'value': 320}
            ]
        else:
            city_distribution = [
                {'city': '北京', 'count': 720},
                {'city': '上海', 'count': 680},
                {'city': '深圳', 'count': 610},
                {'city': '广州', 'count': 520},
                {'city': '杭州', 'count': 470},
                {'city': '成都', 'count': 410}
            ]
            salary_by_city = [
                {'city': '北京', 'avg': 19600.0},
                {'city': '上海', 'avg': 19100.0},
                {'city': '深圳', 'avg': 18600.0},
                {'city': '广州', 'avg': 17100.0},
                {'city': '杭州', 'avg': 16700.0},
                {'city': '成都', 'avg': 15300.0}
            ]
            salary_quantiles = [
                {'city': '北京', 'p25': 13800.0, 'p50': 19200.0, 'p75': 24600.0},
                {'city': '上海', 'p25': 13600.0, 'p50': 18700.0, 'p75': 23900.0},
                {'city': '深圳', 'p25': 13200.0, 'p50': 18300.0, 'p75': 23400.0},
                {'city': '广州', 'p25': 12100.0, 'p50': 16800.0, 'p75': 21600.0},
                {'city': '杭州', 'p25': 11800.0, 'p50': 16400.0, 'p75': 21100.0},
                {'city': '成都', 'p25': 10900.0, 'p50': 15000.0, 'p75': 19800.0}
            ]
            education_distribution = [
                {'name': '本科', 'value': 1550},
                {'name': '大专', 'value': 900},
                {'name': '硕士', 'value': 620},
                {'name': '不限', 'value': 340}
            ]
            experience_distribution = [
                {'name': '不限', 'value': 500},
                {'name': '应届生', 'value': 420},
                {'name': '1-3 年', 'value': 1260},
                {'name': '3-5 年', 'value': 890},
                {'name': '5-10 年', 'value': 340}
            ]
            skill_cloud = [
                {'name': 'Python', 'value': 680},
                {'name': 'Java', 'value': 640},
                {'name': 'SQL', 'value': 620},
                {'name': '数据分析', 'value': 580},
                {'name': '项目管理', 'value': 560},
                {'name': 'Excel', 'value': 540},
                {'name': '机器学习', 'value': 510},
                {'name': 'PowerBI', 'value': 470},
                {'name': '算法', 'value': 440},
                {'name': '产品规划', 'value': 410}
            ]

        return {
            'total_jobs': sum(i['count'] for i in city_distribution),
            'city_distribution': city_distribution,
            'salary_by_city': salary_by_city,
            'salary_quantiles': salary_quantiles,
            'education_distribution': education_distribution,
            'experience_distribution': experience_distribution,
            'skill_cloud': skill_cloud,
            'is_mock': True
        }

    try:
        with get_db_session() as db_session:
            # 使用参数化查询，仅选择必要字段减少内存占用
            search_pattern = f'%{keyword_text}%'
            jobs = db_session.query(
                JobPosition.city,
                JobPosition.education,
                JobPosition.experience,
                JobPosition.min_salary,
                JobPosition.max_salary,
                JobPosition.skill_tags
            ).filter(
                JobPosition.status == 'active',
                (
                    JobPosition.job_name.like(search_pattern) |
                    JobPosition.skill_tags.like(search_pattern) |
                    JobPosition.description.like(search_pattern)
                )
            ).limit(5000).all()

        city_counter = Counter()
        city_salaries = defaultdict(list)
        education_counter = Counter()
        experience_counter = Counter()
        skill_counter = Counter()

        for job in jobs:
            # 解包元组结果
            city_raw, education_raw, experience_raw, min_sal, max_sal, skill_tags = job
            
            city = (city_raw or '未知').strip() or '未知'
            education = (education_raw or '不限').strip() or '不限'
            experience = (experience_raw or '不限').strip() or '不限'

            city_counter[city] += 1
            education_counter[education] += 1
            experience_counter[experience] += 1

            salary = avg_salary(min_sal, max_sal)
            if salary is not None:
                city_salaries[city].append(salary)

            # 严格以技能字段为主进行热词统计
            skill_text = (skill_tags or '').strip()
            if not skill_text:
                continue

            parts = re.split(r'[，,、；;|/\n\t]+', skill_text)
            for token in parts:
                term = token.strip()
                if not term or len(term) < 2 or len(term) > 40:
                    continue
                if term in config.STOP_WORDS:
                    continue
                if re.fullmatch(r'[0-9.\-]+', term):
                    continue

                lower_term = term.lower()
                if term in blocked_exact or lower_term in blocked_exact:
                    continue
                if any(block in term for block in blocked_contains) or any(block in lower_term for block in blocked_contains):
                    continue

                skill_counter[term] += 1

        if len(jobs) == 0:
            mock_payload = build_mock_payload(profession)
            return jsonify({
                'success': True,
                'profession': {
                    'key': profession,
                    'name': keyword_text
                },
                **mock_payload
            })

        top_city_counts = city_counter.most_common(12)
        city_distribution = [
            {'city': city, 'count': count}
            for city, count in top_city_counts
        ]

        salary_by_city = []
        salary_quantiles = []
        for city, _ in top_city_counts:
            values = city_salaries.get(city, [])
            if not values:
                continue

            salary_by_city.append({
                'city': city,
                'avg': round(sum(values) / len(values), 2)
            })

            if len(values) >= 8:
                salary_quantiles.append({
                    'city': city,
                    'p25': round(percentile(values, 0.25), 2),
                    'p50': round(percentile(values, 0.50), 2),
                    'p75': round(percentile(values, 0.75), 2)
                })

        education_distribution = [
            {'name': name, 'value': value}
            for name, value in education_counter.most_common()
        ]
        experience_distribution = [
            {'name': name, 'value': value}
            for name, value in experience_counter.most_common()
        ]
        skill_cloud = [
            {'name': name, 'value': value}
            for name, value in skill_counter.most_common(80)
        ]

        response_data = {
            'profession': {
                'key': profession,
                'name': keyword_text
            },
            'total_jobs': len(jobs),
            'city_distribution': city_distribution,
            'salary_by_city': salary_by_city,
            'salary_quantiles': salary_quantiles,
            'education_distribution': education_distribution,
            'experience_distribution': experience_distribution,
            'skill_cloud': skill_cloud,
            'is_mock': False
        }
        
        # 存入缓存
        INSIGHTS_CACHE['profession_data'][profession] = {
            'expires_at': datetime.now() + timedelta(seconds=INSIGHTS_CACHE_TTL_SECONDS),
            'data': response_data
        }

        return jsonify({'success': True, **response_data})
    except Exception as e:
        app.logger.warning("加载专业洞察数据失败：%s", e)
        mock_payload = build_mock_payload(profession)
        return jsonify({
            'success': True,
            'profession': {
                'key': profession,
                'name': keyword_text
            },
            **mock_payload
        })


@app.route('/api/insights/overall-skill-cloud')
def insights_overall_skill_cloud():
    """全部专业视角：返回全局技能词云数据（动态计算）。"""

    # 检查缓存
    now = datetime.now()
    cache_entry = INSIGHTS_CACHE['overall_skill_cloud']
    if cache_entry['expires_at'] and now < cache_entry['expires_at']:
        app.logger.info("命中全局技能词云缓存")
        return jsonify({'success': True, 'skill_cloud': cache_entry['data'], 'cached': True})

    blocked_exact = {
        '薪资', '经验', '学历', '待遇', '要求', '应届生', '大专', '硕士', '本科', '以下',
        '运营', '人力', '市场', '专员', '分析师', '岗位', '职位', '工作', '相关', '优先', '不限'
    }
    blocked_contains = {
        '薪资', '经验', '学历', '待遇', '应届', '大专', '硕士', '本科', '以下',
        '专员', '分析师', '运营', '人力', '市场', '招聘', '岗位', '职位'
    }

    fallback_cloud = [
        {'name': 'Python', 'value': 680},
        {'name': 'Java', 'value': 640},
        {'name': 'SQL', 'value': 620},
        {'name': '数据分析', 'value': 580},
        {'name': '项目管理', 'value': 560},
        {'name': 'Excel', 'value': 540},
        {'name': '机器学习', 'value': 510},
        {'name': 'PowerBI', 'value': 470},
        {'name': '算法', 'value': 440},
        {'name': '产品规划', 'value': 410}
    ]

    try:
        with get_db_session() as db_session:
            # 仅查询skill_tags字段，减少内存占用
            jobs = db_session.query(JobPosition.skill_tags).filter(
                JobPosition.status == 'active'
            ).limit(8000).all()

        skill_counter = Counter()

        for (skill_tags,) in jobs:  # 解包元组
            skill_text = (skill_tags or '').strip()
            if not skill_text:
                continue

            parts = re.split(r'[，,、；;|/\n\t]+', skill_text)
            for token in parts:
                term = token.strip()
                if not term or len(term) < 2 or len(term) > 40:
                    continue
                if term in config.STOP_WORDS:
                    continue
                if re.fullmatch(r'[0-9.\-]+', term):
                    continue

                lower_term = term.lower()
                if term in blocked_exact or lower_term in blocked_exact:
                    continue
                if any(block in term for block in blocked_contains) or any(block in lower_term for block in blocked_contains):
                    continue

                skill_counter[term] += 1

        cloud = [
            {'name': name, 'value': value}
            for name, value in skill_counter.most_common(100)
        ]

        if not cloud:
            cloud = fallback_cloud

        # 存入缓存
        INSIGHTS_CACHE['overall_skill_cloud'] = {
            'expires_at': datetime.now() + timedelta(seconds=INSIGHTS_CACHE_TTL_SECONDS),
            'data': cloud
        }

        return jsonify({'success': True, 'skill_cloud': cloud})
    except Exception as e:
        app.logger.warning("加载全局技能词云失败：%s", e)
        return jsonify({'success': True, 'skill_cloud': fallback_cloud})


@app.route('/api/insights/overall-data')
def insights_overall_data():
    """全部专业视角：返回省份岗位分布与经验要求分布。"""

    # 检查缓存
    now = datetime.now()
    cache_entry = INSIGHTS_CACHE['overall_data']
    if cache_entry['expires_at'] and now < cache_entry['expires_at']:
        app.logger.info("命中全局数据洞察缓存")
        return jsonify({'success': True, **cache_entry['data'], 'cached': True})

    city_province_map = {
        '北京': '北京',
        '上海': '上海',
        '天津': '天津',
        '重庆': '重庆',
        '广州': '广东',
        '深圳': '广东',
        '佛山': '广东',
        '东莞': '广东',
        '珠海': '广东',
        '杭州': '浙江',
        '宁波': '浙江',
        '温州': '浙江',
        '南京': '江苏',
        '苏州': '江苏',
        '无锡': '江苏',
        '常州': '江苏',
        '成都': '四川',
        '武汉': '湖北',
        '西安': '陕西',
        '长沙': '湖南',
        '郑州': '河南',
        '合肥': '安徽',
        '青岛': '山东',
        '济南': '山东',
        '厦门': '福建',
        '福州': '福建',
        '南昌': '江西',
        '沈阳': '辽宁',
        '大连': '辽宁',
        '长春': '吉林',
        '哈尔滨': '黑龙江'
    }

    fallback_payload = {
        'province_distribution': [
            {'name': '广东', 'count': 920},
            {'name': '北京', 'count': 860},
            {'name': '上海', 'count': 830},
            {'name': '浙江', 'count': 780},
            {'name': '江苏', 'count': 740},
            {'name': '四川', 'count': 510},
            {'name': '湖北', 'count': 430},
            {'name': '陕西', 'count': 390}
        ],
        'city_distribution': [
            {'city': '北京', 'count': 720},
            {'city': '上海', 'count': 680},
            {'city': '深圳', 'count': 610},
            {'city': '广州', 'count': 520},
            {'city': '杭州', 'count': 470},
            {'city': '成都', 'count': 410}
        ],
        'salary_by_city': [
            {'city': '北京', 'avg': 19600.0},
            {'city': '上海', 'avg': 19100.0},
            {'city': '深圳', 'avg': 18600.0},
            {'city': '广州', 'avg': 17100.0},
            {'city': '杭州', 'avg': 16700.0},
            {'city': '成都', 'avg': 15300.0}
        ],
        'salary_quantiles': [
            {'city': '北京', 'p25': 13800.0, 'p50': 19200.0, 'p75': 24600.0},
            {'city': '上海', 'p25': 13600.0, 'p50': 18700.0, 'p75': 23900.0},
            {'city': '深圳', 'p25': 13200.0, 'p50': 18300.0, 'p75': 23400.0},
            {'city': '广州', 'p25': 12100.0, 'p50': 16800.0, 'p75': 21600.0},
            {'city': '杭州', 'p25': 11800.0, 'p50': 16400.0, 'p75': 21100.0},
            {'city': '成都', 'p25': 10900.0, 'p50': 15000.0, 'p75': 19800.0}
        ],
        'education_distribution': [
            {'name': '本科', 'value': 1550},
            {'name': '大专', 'value': 900},
            {'name': '硕士', 'value': 620},
            {'name': '不限', 'value': 340}
        ],
        'experience_distribution': [
            {'name': '不限', 'value': 500},
            {'name': '应届生', 'value': 420},
            {'name': '1-3 年', 'value': 1260},
            {'name': '3-5 年', 'value': 890},
            {'name': '5-10 年', 'value': 340}
        ],
        'is_mock': True
    }

    def avg_salary(min_salary, max_salary):
        nums = [
            float(v)
            for v in [min_salary, max_salary]
            if isinstance(v, (int, float)) and v > 0
        ]
        if not nums:
            return None
        return sum(nums) / len(nums)

    def percentile(values, p):
        if not values: return 0.0
        import statistics
        return float(statistics.quantiles(values, n=100, method='inclusive')[int(p * 100) - 1]) if 0 < p < 1 else float(max(values) if p == 1 else min(values))

    def normalize_province(city_text):
        city = (city_text or '').strip()
        if not city:
            return '未知'
        for city_name, province_name in city_province_map.items():
            if city_name in city:
                return province_name

        if city in {'北京', '上海', '天津', '重庆'}:
            return city
        if city in {'北京市', '上海市', '天津市', '重庆市'}:
            return city[:2]
        if city.endswith('省'):
            return city[:-1]
        if city.endswith('市') and len(city) <= 4:
            return city[:-1]
        return city

    try:
        with get_db_session() as db_session:
            # 仅查询必要字段，减少内存占用
            jobs = db_session.query(
                JobPosition.city,
                JobPosition.experience,
                JobPosition.education,
                JobPosition.min_salary,
                JobPosition.max_salary
            ).filter(
                JobPosition.status == 'active'
            ).limit(10000).all()

        province_counter = Counter()
        city_counter = Counter()
        city_salaries = defaultdict(list)
        experience_counter = Counter()
        education_counter = Counter()

        for job in jobs:
            # 解包元组结果
            city_raw, exp_raw, edu_raw, min_sal, max_sal = job
            
            city = (city_raw or '').strip() or '未知'
            province = normalize_province(city)
            experience = (exp_raw or '不限').strip() or '不限'
            education = (edu_raw or '不限').strip() or '不限'

            province_counter[province] += 1
            city_counter[city] += 1
            experience_counter[experience] += 1
            education_counter[education] += 1

            salary = avg_salary(min_sal, max_sal)
            if salary is not None:
                city_salaries[city].append(salary)

        province_distribution = [
            {'name': name, 'count': count}
            for name, count in province_counter.most_common(15)
        ]
        city_distribution = [
            {'city': city, 'count': count}
            for city, count in city_counter.most_common(12)
        ]
        salary_by_city = []
        salary_quantiles = []

        for city, _ in city_counter.most_common(12):
            values = city_salaries.get(city, [])
            if not values:
                continue

            salary_by_city.append({
                'city': city,
                'avg': round(sum(values) / len(values), 2)
            })

            if len(values) >= 8:
                salary_quantiles.append({
                    'city': city,
                    'p25': round(percentile(values, 0.25), 2),
                    'p50': round(percentile(values, 0.50), 2),
                    'p75': round(percentile(values, 0.75), 2)
                })

        education_distribution = [
            {'name': name, 'value': value}
            for name, value in education_counter.most_common()
        ]
        experience_distribution = [
            {'name': name, 'value': value}
            for name, value in experience_counter.most_common()
        ]

        if not province_distribution or not experience_distribution:
            return jsonify({'success': True, **fallback_payload})

        response_data = {
            'province_distribution': province_distribution,
            'city_distribution': city_distribution,
            'salary_by_city': salary_by_city,
            'salary_quantiles': salary_quantiles,
            'education_distribution': education_distribution,
            'experience_distribution': experience_distribution,
            'is_mock': False
        }
        
        # 存入缓存
        INSIGHTS_CACHE['overall_data'] = {
            'expires_at': datetime.now() + timedelta(seconds=INSIGHTS_CACHE_TTL_SECONDS),
            'data': response_data
        }

        return jsonify({'success': True, **response_data})
    except Exception as e:
        app.logger.warning("加载全部专业聚合数据失败：%s", e)
        return jsonify({'success': True, **fallback_payload})


@app.route('/jobs/search')
def job_search():
    """职位搜索页面"""
    # 获取搜索参数
    keyword = request.args.get('keyword', '').strip()
    city = request.args.get('city', '').strip()
    profession = request.args.get('profession', '').strip()
    min_salary = request.args.get('min_salary', type=int)
    max_salary = request.args.get('max_salary', type=int)
    
    # 如果有搜索参数，直接查询结果
    jobs = []
    if keyword or city or profession or min_salary or max_salary:
        try:
            with get_db_session() as db_session:
                # 构建查询条件
                conditions = []

                # 关键词搜索（使用参数化查询防止SQL注入）
                if keyword:
                    search_pattern = f'%{keyword}%'
                    conditions.append(
                        (JobPosition.job_name.like(search_pattern)) |
                        (JobPosition.company_name.like(search_pattern)) |
                        (JobPosition.description.like(search_pattern)) |
                        (JobPosition.skill_tags.like(search_pattern))
                    )

                # 城市筛选
                if city:
                    actual_city = get_actual_city(city)
                    conditions.append(JobPosition.city == actual_city)

                # 专业方向筛选（通过技能标签）
                if profession:
                    keyword_text = get_profession_keyword(profession)
                    if keyword_text:
                        search_pattern = f'%{keyword_text}%'
                        conditions.append(
                            (JobPosition.skill_tags.like(search_pattern)) |
                            (JobPosition.job_name.like(search_pattern))
                        )

                # 薪资范围
                if min_salary:
                    conditions.append(JobPosition.min_salary >= min_salary)
                if max_salary:
                    conditions.append(JobPosition.max_salary <= max_salary)

                # 查询职位
                query = db_session.query(JobPosition).filter(JobPosition.status == 'active')
                if conditions:
                    query = query.filter(and_(*conditions))

                jobs = query.order_by(JobPosition.id.desc()).limit(100).all()

                # 转换为字典
                jobs = [{
                    'id': job.id,
                    'job_name': job.job_name,
                    'company_name': job.company_name,
                    'city': job.city,
                    'education': job.education,
                    'experience': job.experience,
                    'min_salary': job.min_salary,
                    'max_salary': job.max_salary,
                    'skill_tags': job.skill_tags,
                } for job in jobs]
        except Exception as e:
            app.logger.warning("获取职位失败：%s", e)
    
    return render_template('job_search.html', jobs=jobs)


@app.route('/job/<int:job_id>')
def job_detail(job_id):
    """职位详情页面"""
    return render_template('job_detail.html', job_id=job_id)


@app.route('/api/jobs/search')
def search_jobs():
    """职位搜索 API - 支持多条件筛选"""
    keyword = request.args.get('keyword', '').strip()
    city = request.args.get('city', '').strip()
    profession = request.args.get('profession', '').strip()
    education = request.args.get('education', '').strip()
    experience = request.args.get('experience', '').strip()
    min_salary = request.args.get('min_salary', type=int)
    max_salary = request.args.get('max_salary', type=int)

    if min_salary and max_salary and min_salary > max_salary:
        return jsonify({'success': False, 'message': '最低薪资不能高于最高薪资'}), 400
    
    try:
        with get_db_session() as db_session:
            # 构建查询条件
            conditions = []

            # 关键词搜索（职位名、公司名、描述、技能）- 使用参数化查询
            if keyword:
                search_pattern = f'%{keyword}%'
                conditions.append(
                    (JobPosition.job_name.like(search_pattern)) |
                    (JobPosition.company_name.like(search_pattern)) |
                    (JobPosition.description.like(search_pattern)) |
                    (JobPosition.skill_tags.like(search_pattern))
                )

            # 城市筛选
            if city:
                actual_city = get_actual_city(city)
                conditions.append(JobPosition.city == actual_city)

            # 专业方向筛选（通过技能标签）
            if profession:
                keyword_text = get_profession_keyword(profession)
                if keyword_text:
                    search_pattern = f'%{keyword_text}%'
                    conditions.append(
                        (JobPosition.skill_tags.like(search_pattern)) |
                        (JobPosition.job_name.like(search_pattern))
                    )

            # 学历要求
            if education:
                conditions.append(JobPosition.education == education)

            # 经验要求
            if experience:
                conditions.append(JobPosition.experience == experience)

            # 薪资范围
            if min_salary:
                conditions.append(JobPosition.min_salary >= min_salary)
            if max_salary:
                conditions.append(JobPosition.max_salary <= max_salary)

            # 查询职位
            query = db_session.query(JobPosition).filter(JobPosition.status == 'active')
            if conditions:
                query = query.filter(and_(*conditions))

            jobs = query.order_by(JobPosition.id.desc()).limit(100).all()

            result = [{
                'id': job.id,
                'job_name': job.job_name,
                'company_name': job.company_name,
                'city': job.city,
                'education': job.education,
                'experience': job.experience,
                'min_salary': job.min_salary,
                'max_salary': job.max_salary,
                'skill_keywords': job.skill_tags,
            } for job in jobs]

        return jsonify({'success': True, 'jobs': result, 'total': len(result)})
        
    except Exception:
        app.logger.exception("搜索职位失败")
        return jsonify({'success': False, 'message': '搜索失败，请稍后重试'}), 500


@app.route('/api/jobs/<int:job_id>')
def get_job_detail(job_id):
    """获取单个职位详情 API"""
    try:
        with get_db_session() as db_session:
            job = db_session.query(JobPosition).filter_by(id=job_id).first()

            if not job:
                return jsonify({'success': False, 'message': '职位不存在'}), 404

            # 非活跃职位默认不对普通访问者暴露，仅管理员/企业模式可查看
            if job.status != 'active' and not (session.get('is_admin') or session.get('is_company')):
                return jsonify({'success': False, 'message': '职位不存在'}), 404

            result = {
                'id': job.id,
                'job_name': job.job_name,
                'company_name': job.company_name,
                'min_salary': job.min_salary,
                'max_salary': job.max_salary,
                'salary_text': job.salary_text,
                'city': job.city,
                'district': job.district,
                'experience': job.experience,
                'education': job.education,
                'skill_tags': job.skill_tags,
                'description': job.description,
                'industry': job.industry,
                'company_size': job.company_size,
                'source': job.source,
                'is_campus': bool(job.is_campus),
                'publish_date': job.publish_date,
                'crawl_time': job.crawl_time.strftime('%Y-%m-%d %H:%M:%S') if job.crawl_time else None
            }

        return jsonify({'success': True, 'job': result})
        
    except Exception:
        app.logger.exception("获取职位详情失败")
        return jsonify({'success': False, 'message': '获取职位详情失败，请稍后重试'}), 500


@app.route('/applications')
def applications():
    """投递记录页面——仅普通用户访问"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_active:
            session.clear()
            return redirect('/login?message=登录已失效，请重新登录')

        # 角色强约束：企业管理员和超级管理员不能进入普通用户页
        if user.is_admin:
            return redirect('/admin')
        if user.is_company_admin:
            return redirect('/company')
    
    return render_template('applications.html')


@app.route('/profile')
def profile():
    """个人中心页面——仅普通用户访问"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_active:
            session.clear()
            return redirect('/login?message=登录已失效，请重新登录')

        # 角色强约束：企业管理员和超级管理员不能进入普通用户页
        if user.is_admin:
            return redirect('/admin')
        if user.is_company_admin:
            return redirect('/company')
        
        if user:
            with get_db_session() as resume_session:
                # 获取投递记录数量
                applications_count = resume_session.query(JobApplication).filter_by(
                    user_id=session['user_id']
                ).count()
                
                return render_template('profile.html', 
                                       user=user,
                                       applications_count=applications_count)
    
    return redirect(url_for('login'))


@app.route('/api/account/avatar/<filename>', methods=['GET'])
def get_account_avatar(filename):
    """获取用户头像文件。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    safe_name = secure_filename(filename)
    if not safe_name:
        return jsonify({'success': False, 'message': '无效文件名'}), 400

    basedir = os.path.abspath(os.path.dirname(__file__))
    avatar_dir = os.path.join(basedir, 'uploads', 'avatars')
    file_path = os.path.join(avatar_dir, safe_name)

    if not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '头像不存在'}), 404

    return send_file(file_path)


@app.route('/api/account/avatar', methods=['POST'])
def upload_account_avatar():
    """上传/修改普通用户头像。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403

    if 'avatar' not in request.files:
        return jsonify({'success': False, 'message': '未选择头像文件'}), 400

    file = request.files['avatar']
    if not file or file.filename == '':
        return jsonify({'success': False, 'message': '未选择头像文件'}), 400

    # 检查文件大小（最大 2MB）
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    if file_size > 2 * 1024 * 1024:
        return jsonify({'success': False, 'message': '头像大小不能超过 2MB'}), 400

    original_filename = file.filename
    if '.' not in original_filename:
        return jsonify({'success': False, 'message': '无效的文件名'}), 400

    ext = original_filename.rsplit('.', 1)[1].lower()
    allowed_extensions = {'jpg', 'jpeg', 'png', 'webp'}
    if ext not in allowed_extensions:
        return jsonify({'success': False, 'message': '仅支持 JPG、PNG、WEBP 格式'}), 400

    # 校验文件头，避免伪造扩展名
    header = file.read(16)
    file.seek(0)
    is_jpeg = header.startswith(b'\xff\xd8\xff')
    is_png = header.startswith(b'\x89PNG\r\n\x1a\n')
    is_webp = len(header) >= 12 and header.startswith(b'RIFF') and header[8:12] == b'WEBP'
    if ext in {'jpg', 'jpeg'} and not is_jpeg:
        return jsonify({'success': False, 'message': '无效的 JPG 文件'}), 400
    if ext == 'png' and not is_png:
        return jsonify({'success': False, 'message': '无效的 PNG 文件'}), 400
    if ext == 'webp' and not is_webp:
        return jsonify({'success': False, 'message': '无效的 WEBP 文件'}), 400

    basedir = os.path.abspath(os.path.dirname(__file__))
    avatar_dir = os.path.join(basedir, 'uploads', 'avatars')
    os.makedirs(avatar_dir, exist_ok=True)

    import uuid
    safe_original = secure_filename(original_filename)
    if not safe_original:
        safe_original = f"avatar.{ext}"
    unique_filename = f"u{session['user_id']}_{uuid.uuid4().hex}_{safe_original}"
    file_path = os.path.join(avatar_dir, unique_filename)
    file.save(file_path)

    avatar_url = f"/api/account/avatar/{unique_filename}"

    user_session = UserSession()
    try:
        user = user_session.query(User).filter_by(id=session['user_id']).first()
        if not user:
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'success': False, 'message': '用户不存在'}), 404

        old_avatar = user.avatar
        user.avatar = avatar_url
        user_session.commit()

        # 若旧头像在本地 uploads/avatars 下，替换后删除旧文件
        if old_avatar and old_avatar.startswith('/api/account/avatar/'):
            old_filename = secure_filename(old_avatar.rsplit('/', 1)[-1])
            old_path = os.path.join(avatar_dir, old_filename)
            if old_filename and old_filename != unique_filename and os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except Exception as cleanup_error:
                    app.logger.warning("删除旧头像失败: %s", cleanup_error)

        app.logger.info(
            f"头像上传: User ID {session['user_id']} "
            f"file={original_filename} size={file_size}bytes from {get_client_ip()}"
        )
        return jsonify({'success': True, 'message': '头像上传成功', 'avatar_url': avatar_url})
    except Exception:
        user_session.rollback()
        app.logger.exception("头像上传失败")
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
        return jsonify({'success': False, 'message': '头像上传失败，请重试'}), 500
    finally:
        user_session.close()


@app.route('/api/account/avatar', methods=['DELETE'])
def delete_account_avatar():
    """恢复为默认头像（删除当前头像）。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403

    basedir = os.path.abspath(os.path.dirname(__file__))
    avatar_dir = os.path.join(basedir, 'uploads', 'avatars')

    user_session = UserSession()
    try:
        user = user_session.query(User).filter_by(id=session['user_id']).first()
        if not user:
            return jsonify({'success': False, 'message': '用户不存在'}), 404

        old_avatar = user.avatar
        user.avatar = None
        user_session.commit()

        if old_avatar and old_avatar.startswith('/api/account/avatar/'):
            old_filename = secure_filename(old_avatar.rsplit('/', 1)[-1])
            old_path = os.path.join(avatar_dir, old_filename)
            if old_filename and os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except Exception as cleanup_error:
                    app.logger.warning("删除头像文件失败: %s", cleanup_error)

        app.logger.info(
            f"头像删除: User ID {session['user_id']} from {get_client_ip()}"
        )
        return jsonify({'success': True, 'message': '已恢复默认头像'})
    except Exception:
        user_session.rollback()
        app.logger.exception("删除头像失败")
        return jsonify({'success': False, 'message': '恢复默认头像失败，请重试'}), 500
    finally:
        user_session.close()


@app.route('/admin')
def admin_panel():
    """管理员后台"""
    if 'user_id' not in session:
        return redirect('/login')
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return redirect('/login')
        
        return render_template('admin.html')


@app.route('/api/account/settings', methods=['POST'])
def account_settings():
    """账户设置 API"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    data = request.get_json(silent=True) or {}
    username = data.get('username', '').strip()
    email = data.get('email', '').strip()
    phone = data.get('phone', '').strip()
    graduation_year = data.get('graduation_year')
    education = data.get('education', '')
    current_password = data.get('current_password', '')
    new_password = data.get('new_password', '')
    
    # 验证必填项
    if not username or not email:
        return jsonify({'success': False, 'message': '用户名和邮箱为必填项'}), 400
    
    if len(username) < 3 or len(username) > 20:
        return jsonify({'success': False, 'message': '用户名长度必须在 3-20 个字符之间'}), 400
    
    from werkzeug.security import check_password_hash, generate_password_hash
    
    user_session = UserSession()
    try:
        user = user_session.query(User).filter_by(id=session['user_id']).first()
        if not user:
            return jsonify({'success': False, 'message': '用户不存在'}), 404
        
        # 检查用户名是否已被其他用户使用
        existing_user = user_session.query(User).filter(
            User.username == username,
            User.id != session['user_id']
        ).first()
        if existing_user:
            return jsonify({'success': False, 'message': '用户名已被使用'}), 400
        
        # 检查邮箱是否已被其他用户使用
        existing_email = user_session.query(User).filter(
            User.email == email,
            User.id != session['user_id']
        ).first()
        if existing_email:
            return jsonify({'success': False, 'message': '邮箱已被使用'}), 400
        
        # 如果要修改密码，验证当前密码
        if current_password and new_password:
            if not check_password_hash(user.password_hash, current_password):
                return jsonify({'success': False, 'message': '当前密码错误'}), 400
            
            if len(new_password) < 8:
                return jsonify({'success': False, 'message': '新密码长度至少 8 位'}), 400
            
            # 更新密码
            user.password_hash = generate_password_hash(new_password)

        parsed_graduation_year = None
        if graduation_year not in (None, ''):
            try:
                parsed_graduation_year = int(graduation_year)
            except (TypeError, ValueError):
                return jsonify({'success': False, 'message': '毕业年份格式无效'}), 400

            if parsed_graduation_year < 1900 or parsed_graduation_year > 2100:
                return jsonify({'success': False, 'message': '毕业年份超出有效范围'}), 400
        
        # 更新用户信息
        user.username = username
        user.email = email
        user.phone = phone if phone else None
        user.graduation_year = parsed_graduation_year
        user.education = education if education else None
        
        user_session.commit()
        
        # 更新 session 中的用户名
        session['username'] = user.username
        session['email'] = user.email
        
        # 记录账户修改审计日志
        changes = []
        if current_password and new_password:
            changes.append('password')
        app.logger.info(
            f"账户信息修改: User ID {session['user_id']} "
            f"changes={','.join(changes) if changes else 'profile'} "
            f"from {get_client_ip()}"
        )
        
        return jsonify({'success': True, 'message': '保存成功'})
        
    except Exception:
        app.logger.exception("保存账户设置失败")
        return jsonify({'success': False, 'message': '保存失败，请重试'}), 500
    finally:
        user_session.close()


# ==================== 简历相关 API ====================

@app.route('/api/resume', methods=['GET'])
def get_resume():
    """获取用户简历——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        if resume:
            return jsonify({'success': True, 'resume': resume.to_dict()})
        else:
            return jsonify({'success': True, 'resume': None})
    except Exception:
        app.logger.exception("获取简历失败")
        return jsonify({'success': False, 'message': '获取简历失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/resume/attachment/analyze', methods=['POST'])
def analyze_uploaded_resume():
    """接收 multipart/form-data 的简历文件（或使用已保存附件），提取文本并返回结构化分析。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    # 仅普通用户可访问
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403

    file = request.files.get('file')
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()

        # 如果请求中没有文件，则尝试使用已保存的附件
        if not file:
            if not resume or not resume.attachment_path or not os.path.exists(resume.attachment_path):
                return jsonify({'success': False, 'message': '没有可分析的简历文件'}), 400
            filepath = resume.attachment_path
            original_filename = resume.attachment_name or os.path.basename(filepath)
            ext = (resume.attachment_type or '').lower()
        else:
            # 临时保存上传文件用于解析
            original_filename = file.filename or 'uploaded_resume'
            if '.' in original_filename:
                ext = original_filename.rsplit('.', 1)[1].lower()
            else:
                ext = ''

            # 简单校验
            if ext not in ('pdf', 'doc', 'docx', 'png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif'):
                return jsonify({'success': False, 'message': '仅支持 PDF、Word 和图片格式'}), 400

            basedir = os.path.abspath(os.path.dirname(__file__))
            temp_dir = os.path.join(basedir, 'uploads', 'resumes_temp')
            os.makedirs(temp_dir, exist_ok=True)
            import uuid
            temp_name = f"{uuid.uuid4().hex}_{secure_filename(original_filename)}"
            filepath = os.path.join(temp_dir, temp_name)
            file.save(filepath)

        # 对图片尝试检测 OCR 可用性，供前端展示提示
        ocr_attempted = False
        ocr_available = False
        if ext in ('png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif'):
            ocr_attempted = True
            try:
                import pytesseract  # type: ignore
                ocr_available = True
            except Exception:
                ocr_available = False

        # 提取文本
        source_text = extract_text_from_file(filepath, ext)
        attachment_text = source_text or ''
        attachment_text_truncated = False
        if len(attachment_text) > 4000:
            attachment_text = attachment_text[:4000]
            attachment_text_truncated = True

        # 如果提取为空，则直接使用本地兜底分析
        template_key = str(request.form.get('template') or request.args.get('template') or 'general')

        # 构建后端分析调用：复用 analyze_resume_workbench 的策略
        resume_data = resume.to_dict() if resume else {}
        fallback_analysis = build_resume_ai_fallback_analysis(source_text, original_filename, template_key, resume_data)

        ai_config = getattr(config, 'AI_CONFIG', {})
        api_key = str(ai_config.get('API_KEY', '')).strip()
        base_url = str(ai_config.get('BASE_URL', '')).strip().rstrip('/')
        model = str(ai_config.get('MODEL', 'qwen-plus')).strip()

        if not api_key or not base_url:
            return jsonify({'success': True, 'analysis': fallback_analysis, 'source': 'local', 'model': model, 'resume': resume_data, 'provider_error': 'ai_config_missing', 'ocr_attempted': ocr_attempted, 'ocr_available': ocr_available, 'ocr_text_len': len(source_text or ''), 'attachment_text': attachment_text, 'attachment_text_truncated': attachment_text_truncated})

        system_prompt = (
            '你是一个招聘平台的简历分析助手。'
            '请基于用户简历和补充文本给出结构化分析。'
            '必须只输出 JSON，不要输出多余解释。JSON 字段必须包含：'
            'score(0-100整数), headline, summary, structure, quant, keywords, priority, strengths(字符串数组), issues(字符串数组), rewrite。'
            '其中 structure/quant/keywords/priority 使用简短中文结论。'
        )

        user_prompt = '\n'.join([
            f"模板视角：{template_key}",
            f"文件名：{original_filename or '无'}",
            f"简历信息：{json.dumps(resume_data, ensure_ascii=False)}",
            f"补充文本：{source_text or '无'}",
            '请直接返回 JSON。',
        ])

        attempts = []
        parsed_analysis = None
        used_provider_url = None
        attachment_timeout = int(ai_config.get('ATTACHMENT_TIMEOUT', ai_config.get('TIMEOUT', 30)))
        for candidate in build_ai_request_candidates(base_url, model, user_prompt, ai_config, system_prompt=system_prompt):
            try:
                response = requests.post(
                    candidate['url'],
                    headers={
                        'Authorization': f'Bearer {api_key}',
                        'Content-Type': 'application/json'
                    },
                    json=candidate['payload'],
                    timeout=attachment_timeout
                )
                response.raise_for_status()
                result = response.json()
                reply = extract_ai_reply(result)
                parsed = parse_ai_json_reply(reply)
                if parsed:
                    parsed_analysis = normalize_resume_ai_analysis(parsed, fallback_analysis)
                    used_provider_url = candidate['url']
                    break

                attempts.append(f"{candidate['url']}: 返回成功但未解析到 JSON")
            except requests.HTTPError as http_exc:
                detail = extract_provider_error(getattr(http_exc, 'response', None))
                attempts.append(f"{candidate['url']}: {http_exc}{' - ' + detail if detail else ''}")
                if detail and 'does not support http call' in detail.lower():
                    continue
            except requests.RequestException as exc:
                attempts.append(f"{candidate['url']}: {exc}")
            except ValueError:
                attempts.append(f"{candidate['url']}: 接口返回了非 JSON 内容")

        if parsed_analysis is None:
            app.logger.warning('简历附件分析 AI 不可用，回退到本地分析: %s', attempts)
            return jsonify({'success': True, 'analysis': fallback_analysis, 'source': 'local', 'model': model, 'resume': resume_data, 'provider_error': 'analysis_fallback', 'ocr_attempted': ocr_attempted, 'ocr_available': ocr_available, 'ocr_text_len': len(source_text or ''), 'attachment_text': attachment_text, 'attachment_text_truncated': attachment_text_truncated})

        # 成功返回解析结果
        # 成功返回解析结果
        return jsonify({'success': True, 'analysis': parsed_analysis, 'source': 'ai', 'model': model, 'provider_url': used_provider_url, 'resume': resume_data, 'ocr_attempted': ocr_attempted, 'ocr_available': ocr_available, 'ocr_text_len': len(source_text or ''), 'attachment_text': attachment_text, 'attachment_text_truncated': attachment_text_truncated})

    except Exception:
        app.logger.exception('处理简历附件分析失败')
        return jsonify({'success': True, 'analysis': build_resume_ai_fallback_analysis('', '', 'general'), 'source': 'local', 'provider_error': 'analyze_route_exception', 'ocr_attempted': False, 'ocr_available': False, 'ocr_text_len': 0, 'attachment_text': '', 'attachment_text_truncated': False})
    finally:
        resume_session.close()


@app.route('/api/resume', methods=['POST'])
def save_resume():
    """保存/更新简历——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    data = request.get_json(silent=True) or {}
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume:
            # 创建新简历
            resume = Resume(user_id=session['user_id'])
            resume_session.add(resume)
        
        # 更新简历信息
        resume.name = data.get('name') or ''
        resume.phone = data.get('phone') or ''
        resume.email = data.get('email') or ''
        resume.education = data.get('education') or ''
        resume.graduation_year = data.get('graduation_year') or None
        resume.school = data.get('school') or ''
        resume.major = data.get('major') or ''
        resume.work_experience = data.get('work_experience') or ''
        resume.project_experience = data.get('project_experience') or ''
        resume.skills = data.get('skills') or ''
        resume.self_introduction = data.get('self_introduction') or ''
        resume.expected_city = data.get('expected_city') or ''
        resume.expected_salary = data.get('expected_salary') or ''
        resume.career_preference = data.get('career_preference') or ''
        
        # 检查简历是否完善（只要有基本信息就算完善）
        if resume.name and resume.phone and resume.email:
            resume.is_complete = 1
        
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '简历保存成功', 'resume': resume.to_dict()})
    except Exception:
        resume_session.rollback()
        app.logger.exception("保存简历失败")
        return jsonify({'success': False, 'message': '保存失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/resume/analysis', methods=['POST'])
def analyze_resume_workbench():
    """为投递中心的简历工作台生成后端分析结果。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403

    data = request.get_json(silent=True) or {}
    template_key = str(data.get('template') or 'general').strip() or 'general'
    source_text = str(data.get('source_text') or '').strip()
    file_name = str(data.get('file_name') or '').strip()

    ai_config = getattr(config, 'AI_CONFIG', {})
    api_key = str(ai_config.get('API_KEY', '')).strip()
    base_url = str(ai_config.get('BASE_URL', '')).strip().rstrip('/')
    model = str(ai_config.get('MODEL', 'qwen-plus')).strip()

    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        resume_data = resume.to_dict() if resume else {}
        fallback_analysis = build_resume_ai_fallback_analysis(source_text, file_name, template_key, resume_data)

        if not api_key or not base_url:
            return jsonify({
                'success': True,
                'analysis': fallback_analysis,
                'source': 'local',
                'model': model,
                'resume': resume_data,
                'provider_error': 'ai_config_missing'
            })

        system_prompt = (
            '你是一个招聘平台的简历分析助手。'
            '请基于用户简历和补充文本给出结构化分析。'
            '必须只输出 JSON，不要输出多余解释。JSON 字段必须包含：'
            'score(0-100整数), headline, summary, structure, quant, keywords, priority, strengths(字符串数组), issues(字符串数组), rewrite。'
            '其中 structure/quant/keywords/priority 使用简短中文结论。'
        )

        user_prompt = '\n'.join([
            f"模板视角：{template_key}",
            f"文件名：{file_name or '无'}",
            f"简历信息：{json.dumps(resume_data, ensure_ascii=False)}",
            f"补充文本：{source_text or '无'}",
            '请直接返回 JSON。',
        ])

        attempts = []
        parsed_analysis = None
        used_provider_url = None
        for candidate in build_ai_request_candidates(base_url, model, user_prompt, ai_config, system_prompt=system_prompt):
            try:
                response = requests.post(
                    candidate['url'],
                    headers={
                        'Authorization': f'Bearer {api_key}',
                        'Content-Type': 'application/json'
                    },
                    json=candidate['payload'],
                    timeout=ai_config.get('TIMEOUT', 60)
                )
                response.raise_for_status()
                result = response.json()
                reply = extract_ai_reply(result)
                parsed = parse_ai_json_reply(reply)
                if parsed:
                    parsed_analysis = normalize_resume_ai_analysis(parsed, fallback_analysis)
                    used_provider_url = candidate['url']
                    break

                attempts.append(f"{candidate['url']}: 返回成功但未解析到 JSON")
            except requests.HTTPError as http_exc:
                detail = extract_provider_error(getattr(http_exc, 'response', None))
                attempts.append(f"{candidate['url']}: {http_exc}{' - ' + detail if detail else ''}")
                if detail and 'does not support http call' in detail.lower():
                    continue
            except requests.RequestException as exc:
                attempts.append(f"{candidate['url']}: {exc}")
            except ValueError:
                attempts.append(f"{candidate['url']}: 接口返回了非 JSON 内容")

        if parsed_analysis is None:
            app.logger.warning('简历分析 AI 不可用，回退到本地分析: %s', attempts)
            return jsonify({
                'success': True,
                'analysis': fallback_analysis,
                'source': 'local',
                'model': model,
                'resume': resume_data,
                'provider_error': attempts[-1] if attempts else 'analysis_fallback'
            })

        return jsonify({
            'success': True,
            'analysis': parsed_analysis,
            'source': 'ai',
            'model': model,
            'provider_url': used_provider_url,
            'resume': resume_data
        })
    except Exception:
        app.logger.exception('生成简历分析失败')
        return jsonify({
            'success': True,
            'analysis': build_resume_ai_fallback_analysis(source_text, file_name, template_key),
            'source': 'local',
            'provider_error': 'analysis_route_exception'
        })
    finally:
        resume_session.close()


@app.route('/api/resume/analysis/export', methods=['POST'])
def export_resume_analysis():
    """导出简历分析结果为 Word 文档。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403

    data = request.get_json(silent=True) or {}
    analysis = data.get('analysis') or {}
    if not isinstance(analysis, dict):
        return jsonify({'success': False, 'message': '分析数据无效'}), 400

    file_name = str(data.get('file_name') or '简历').strip() or '简历'
    template_key = str(data.get('template') or '通用').strip() or '通用'

    try:
        buffer = build_resume_ai_docx(analysis, file_name, template_key)
    except RuntimeError:
        return jsonify({'success': False, 'message': '服务端未安装 Word 导出依赖'}), 500
    except Exception:
        app.logger.exception('生成简历分析导出文件失败')
        return jsonify({'success': False, 'message': '导出失败'}), 500

    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    download_name = f"简历优化建议_{timestamp}.docx"
    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=download_name
    )


@app.route('/api/resume/upload', methods=['POST'])
def upload_resume():
    """上传简历附件——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    # 检查文件大小（最大 10MB）
    file.seek(0, 2)  # 移动到文件末尾
    file_size = file.tell()
    file.seek(0)  # 重置文件指针
    
    if file_size > 10 * 1024 * 1024:
        return jsonify({'success': False, 'message': '文件大小不能超过 10MB'}), 400
    
    # 检查文件类型
    allowed_extensions = {'pdf', 'doc', 'docx', 'png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif'}
    original_filename = file.filename
    if '.' not in original_filename:
        return jsonify({'success': False, 'message': '无效的文件名'}), 400
    
    ext = original_filename.rsplit('.', 1)[1].lower()
    
    if ext not in allowed_extensions:
        return jsonify({'success': False, 'message': '仅支持 PDF、Word 和图片格式'}), 400
    
    # 验证文件内容（检查文件头）
    file_bytes = file.read(1024)  # 读取前 1KB
    file.seek(0)  # 重置文件指针
    
    # PDF 文件头：%PDF
    # DOCX 文件头：PK (ZIP 格式)
    # DOC 文件头：\xD0\xCF\x11\xE0
    # 图片文件使用 Pillow 做快速格式校验
    if ext == 'pdf' and not file_bytes.startswith(b'%PDF'):
        return jsonify({'success': False, 'message': '无效的 PDF 文件'}), 400
    elif ext in ['docx'] and not file_bytes.startswith(b'PK'):
        return jsonify({'success': False, 'message': '无效的 Word 文件'}), 400
    elif ext == 'doc' and not file_bytes.startswith(b'\xD0\xCF\x11\xE0'):
        return jsonify({'success': False, 'message': '无效的 Word 文件'}), 400
    elif ext in ('png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif'):
        try:
            from PIL import Image
            from io import BytesIO
            Image.open(BytesIO(file_bytes)).verify()
        except Exception:
            return jsonify({'success': False, 'message': '无效的图片文件'}), 400
    
    # 创建上传目录
    basedir = os.path.abspath(os.path.dirname(__file__))
    upload_dir = os.path.join(basedir, 'uploads', 'resumes')
    os.makedirs(upload_dir, exist_ok=True)
    
    # 保存文件（使用安全文件名）
    import uuid
    safe_original = secure_filename(original_filename)  # 清理原始文件名
    if not safe_original:
        safe_original = 'resume_' + str(uuid.uuid4().hex)[:8] + '.' + ext
    unique_filename = f"{uuid.uuid4().hex}_{safe_original}"
    filepath = os.path.join(upload_dir, unique_filename)
    file.save(filepath)
    
    # 记录文件上传审计日志
    app.logger.info(
        f"简历附件上传: User ID {session['user_id']} "
        f"file={original_filename} size={file_size}bytes "
        f"from {get_client_ip()}"
    )
    
    # 更新简历
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume:
            resume = Resume(user_id=session['user_id'])
            resume_session.add(resume)
        
        resume.has_attachment = 1
        resume.attachment_path = filepath
        resume.attachment_name = original_filename
        resume.attachment_type = ext
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '上传成功', 'filename': original_filename})
    except Exception as e:
        resume_session.rollback()
        app.logger.exception("上传简历失败")
        return jsonify({'success': False, 'message': '上传失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/resume/attachment', methods=['DELETE'])
def delete_resume_attachment():
    """删除简历附件"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume:
            return jsonify({'success': False, 'message': '没有简历可删除'}), 404
        
        # 删除文件
        if resume.attachment_path and os.path.exists(resume.attachment_path):
            try:
                os.remove(resume.attachment_path)
            except Exception as e:
                app.logger.warning(f"删除文件失败：{e}")
        
        # 更新数据库
        resume.has_attachment = 0
        resume.attachment_path = None
        resume.attachment_name = None
        resume.attachment_type = None
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '删除成功'})
    except Exception as e:
        resume_session.rollback()
        app.logger.exception("删除简历失败")
        return jsonify({'success': False, 'message': '删除失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/resume/attachment/preview')
def preview_resume_attachment():
    """预览简历附件"""
    if 'user_id' not in session:
        return redirect('/login')
    
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume or not resume.attachment_path:
            return jsonify({'error': '没有简历'}), 404
        
        # 检查文件是否存在
        if not os.path.exists(resume.attachment_path):
            return jsonify({'error': '文件不存在'}), 404
        
        attachment_type = (resume.attachment_type or '').lower()

        # docx 优先返回可直接内嵌展示的 HTML，降低对前端组件依赖
        if attachment_type == 'docx':
            try:
                html_content = extract_docx_preview_html(resume.attachment_path)
                return Response(html_content, mimetype='text/html; charset=utf-8')
            except Exception:
                app.logger.exception('DOCX 转 HTML 预览失败，回退到原始文件流')

        if attachment_type == 'doc':
            try:
                html_content = extract_doc_preview_html(resume.attachment_path)
                return Response(html_content, mimetype='text/html; charset=utf-8')
            except Exception:
                app.logger.exception('DOC 转 HTML 预览失败，回退到原始文件流')

        # 根据文件类型设置 MIME
        if attachment_type == 'pdf':
            mime_type = 'application/pdf'
        elif attachment_type in ['doc', 'docx']:
            mime_type = 'application/msword' if attachment_type == 'doc' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif attachment_type in ['png', 'jpeg', 'jpg', 'webp', 'bmp', 'gif']:
            mime_type = 'image/jpeg' if attachment_type in ['jpg', 'jpeg'] else f'image/{attachment_type}'
        else:
            mime_type = 'application/octet-stream'

        return send_file(
            resume.attachment_path,
            mimetype=mime_type,
            as_attachment=False,
            download_name=resume.attachment_name or 'resume.pdf'
        )
    except Exception as e:
        app.logger.exception("预览简历失败")
        return jsonify({'error': '预览失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/resume/download/<int:resume_id>')
def download_resume_by_id(resume_id):
    """根据 ID 下载简历（用于企业管理员）"""
    if 'user_id' not in session:
        return jsonify({'error': '请先登录'}), 401

    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_company_admin or not user.company_name:
                return jsonify({'error': '无权限'}), 403

            # 仅允许下载当前企业收到的投递简历
            related_application = db_session.query(JobApplication).filter_by(
                resume_id=resume_id,
                company_name=user.company_name
            ).order_by(JobApplication.applied_at.desc()).first()

            if not related_application:
                return jsonify({'error': '无权限访问该简历'}), 403

            resume = db_session.query(Resume).filter_by(id=resume_id).first()

            if not resume or not resume.attachment_path:
                return jsonify({'error': '简历不存在'}), 404

            # 检查文件是否存在
            if not os.path.exists(resume.attachment_path):
                return jsonify({'error': '文件不存在'}), 404

            return send_file(
                resume.attachment_path,
                as_attachment=True,
                download_name=resume.attachment_name or 'resume.pdf'
            )
    except Exception as e:
        app.logger.exception("下载简历失败（企业管理员）")
        return jsonify({'error': '下载失败'}), 500


@app.route('/api/resume/attachment/download')
def download_resume_attachment():
    """下载简历附件——仅普通用户"""
    if 'user_id' not in session:
        return redirect('/login')
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'error': '无权限：仅普通用户可访问'}), 403
    
    resume_session = ResumeSession()
    try:
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume or not resume.attachment_path:
            return jsonify({'error': '没有简历'}), 404
        
        # 检查文件是否存在
        if not os.path.exists(resume.attachment_path):
            return jsonify({'error': '文件不存在'}), 404
        
        # 优先使用用户上传时的原始文件名
        file_name = resume.attachment_name or os.path.basename(resume.attachment_path)

        return send_file(
            resume.attachment_path,
            as_attachment=True,
            download_name=file_name
        )
    except Exception as e:
        app.logger.exception("下载简历失败（用户附件）")
        return jsonify({'error': '下载失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/application/submit', methods=['POST'])
def submit_application():
    """提交职位申请——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    data = request.get_json(silent=True) or {}
    job_id = data.get('job_id')
    job_name = data.get('job_name', '')
    company_name = data.get('company_name', '')
    expected_salary = data.get('expected_salary', '')
    reason = data.get('reason', '')
    other_info = data.get('other_info', '')
    
    if not job_id:
        return jsonify({'success': False, 'message': '职位信息不完整'}), 400

    try:
        job_id = int(job_id)
        if job_id <= 0:
            raise ValueError
    except (TypeError, ValueError):
        return jsonify({'success': False, 'message': '职位参数无效'}), 400
    
    resume_session = ResumeSession()
    try:
        # 获取用户简历
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        
        if not resume:
            return jsonify({'success': False, 'message': '请先创建简历'}), 400
        
        # 创建申请记录
        application = JobApplication(
            user_id=session['user_id'],
            resume_id=resume.id,
            job_id=job_id,
            job_name=job_name,
            company_name=company_name,
            applicant_name=resume.name or (user.username if user else None),
            applicant_phone=resume.phone or (user.phone if user else None),
            applicant_email=resume.email or (user.email if user else None),
            expected_salary=expected_salary,
            reason=reason,
            other_info=other_info
        )
        
        resume_session.add(application)
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '申请提交成功'})
    except Exception:
        resume_session.rollback()
        app.logger.exception("提交申请失败")
        return jsonify({'success': False, 'message': '提交失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/applications')
def get_applications():
    """获取用户的投递记录——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    resume_session = ResumeSession()
    try:
        applications = resume_session.query(JobApplication).filter_by(
            user_id=session['user_id']
        ).order_by(JobApplication.applied_at.desc()).all()
        
        result = [app.to_dict() for app in applications]
        return jsonify({'success': True, 'applications': result, 'total': len(result)})
    except Exception:
        app.logger.exception("获取投递记录失败")
        return jsonify({'success': False, 'message': '获取失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/application/<int:app_id>')
def get_application_detail(app_id):
    """获取申请详情——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    resume_session = ResumeSession()
    try:
        application = resume_session.query(JobApplication).filter_by(
            id=app_id,
            user_id=session['user_id']
        ).first()
        
        if not application:
            return jsonify({'success': False, 'message': '申请不存在'}), 404
        
        return jsonify({'success': True, 'application': application.to_dict()})
    except Exception:
        app.logger.exception("获取申请详情失败")
        return jsonify({'success': False, 'message': '获取失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/application/<int:app_id>/withdraw', methods=['POST'])
def withdraw_application(app_id):
    """撤回申请——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    resume_session = ResumeSession()
    try:
        application = resume_session.query(JobApplication).filter_by(
            id=app_id,
            user_id=session['user_id']
        ).first()
        
        if not application:
            return jsonify({'success': False, 'message': '申请不存在'}), 404
        
        # 只有已提交或已查看的申请可以撤回
        if application.status not in ['submitted', 'viewed']:
            return jsonify({'success': False, 'message': '当前状态无法撤回'}), 400
        
        # 删除申请记录
        resume_session.delete(application)
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '申请已撤回'})
    except Exception:
        resume_session.rollback()
        app.logger.exception("撤回申请失败")
        return jsonify({'success': False, 'message': '撤回失败'}), 500
    finally:
        resume_session.close()


@app.route('/api/application/<int:app_id>', methods=['PUT'])
def update_application(app_id):
    """更新申请——仅普通用户"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    # 检查是否为管理员或企业 HR
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if user and (user.is_company_admin or user.is_admin):
            return jsonify({'success': False, 'message': '无权限：仅普通用户可访问'}), 403
    
    data = request.get_json(silent=True) or {}
    
    resume_session = ResumeSession()
    try:
        user = resume_session.query(User).filter_by(id=session['user_id']).first()
        application = resume_session.query(JobApplication).filter_by(
            id=app_id,
            user_id=session['user_id']
        ).first()
        
        if not application:
            return jsonify({'success': False, 'message': '申请不存在'}), 404
        
        # 更新申请信息
        if 'expected_salary' in data:
            application.expected_salary = data['expected_salary']
        if 'reason' in data:
            application.reason = data['reason']
        if 'other_info' in data:
            application.other_info = data['other_info']

        # 同步本次申请的申请人快照（只影响当前申请，不影响其他历史申请）
        resume = resume_session.query(Resume).filter_by(user_id=session['user_id']).first()
        if resume:
            application.applicant_name = resume.name or application.applicant_name or (user.username if user else None)
            application.applicant_phone = resume.phone or application.applicant_phone or (user.phone if user else None)
            application.applicant_email = resume.email or application.applicant_email or (user.email if user else None)
        
        resume_session.commit()
        
        return jsonify({'success': True, 'message': '修改成功'})
    except Exception:
        resume_session.rollback()
        app.logger.exception("更新申请失败")
        return jsonify({'success': False, 'message': '修改失败'}), 500
    finally:
        resume_session.close()


@app.route('/company')
def company_panel():
    """企业管理员后台 - 必须勾选"企业登录"才能访问"""
    if 'user_id' not in session:
        return redirect('/login')
    
    # 关键：检查登录时是否选择了企业登录
    is_company_login = session.get('is_company', False)
    
    if not is_company_login:
        # 如果是企业 HR 账号但没勾选企业登录，提示重新登录
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if user and user.is_company_admin:
                return redirect('/login?message=请使用企业登录方式登录')
        return redirect('/login')
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return redirect('/login')
        
        return render_template('company.html')


@app.route('/company/settings')
def company_settings_page():
    """兼容旧入口：回到企业后台内嵌设置页。"""
    if 'user_id' not in session:
        return redirect('/login')
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return redirect('/login')

        return redirect('/company?tab=settings')


@app.route('/api/company/settings')
def get_company_settings():
    """获取公司设置"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403

        user_company_name = (user.company_name or '').strip()
        if not user_company_name:
            return jsonify({'success': False, 'message': '企业账号未绑定公司，请联系管理员'}), 403

        company = find_company_record_by_name(db_session, user_company_name)
        
        if company:
            # 对齐账号绑定公司名，避免后续请求因名称漂移查不到数据。
            if normalize_company_name(user.company_name) != normalize_company_name(company.company_name):
                user.company_name = company.company_name
                db_session.commit()

            company_data = company.to_dict()
            company_data['recruitment_declaration'] = company_data.get('recruitment_manifesto')
            return jsonify({'success': True, 'company': company_data})
        else:
            # 如果没有公司记录，返回默认值
            return jsonify({
                'success': True,
                'company': {
                    'company_name': user_company_name,
                    'english_name': None,
                    'short_name': None,
                    'credit_code': None,
                    'established_date': None,
                    'industry': None,
                    'website': None,
                    'phone': None,
                    'email': None,
                    'address': None,
                    'description': None,
                    'recruitment_manifesto': None,
                    'recruitment_declaration': None,
                    'logo_path': None,
                    'enable_employer_brand': False,
                    'culture_video_url': None,
                    'employee_stories': None,
                    'resume_retention_days': 365,
                    'interview_record_archive_days': 180,
                    'gdpr_compliance': False,
                    'consent_popup_enabled': True,
                    'data_export_enabled': True,
                    'data_deletion_enabled': True,
                    'audit_log_enabled': True,
                    'audit_log_retention_days': 90,
                    'phone_mask_enabled': True,
                    'email_mask_enabled': False
                }
            })


@app.route('/api/company/settings', methods=['PUT'])
def update_company_settings():
    """更新公司设置"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403

        data = request.get_json(silent=True) or {}
        user_company_name = (user.company_name or '').strip()
        requested_company_name = (data.get('company_name') or '').strip() if 'company_name' in data else None

        if 'company_name' in data and not requested_company_name:
            return jsonify({'success': False, 'message': '公司名称不能为空'}), 400

        lookup_name = user_company_name or requested_company_name
        if not lookup_name:
            return jsonify({'success': False, 'message': '企业账号未绑定公司，请联系管理员'}), 403
        
        # 查找或创建公司记录
        company = find_company_record_by_name(db_session, lookup_name)
        if not company:
            company = Company(company_name=lookup_name)
            db_session.add(company)
        
        # 更新字段
        if requested_company_name:
            conflict_company = db_session.query(Company).filter(
                Company.company_name == requested_company_name,
                Company.id != company.id
            ).first()
            if conflict_company:
                return jsonify({'success': False, 'message': '公司名称已存在，请检查后重试'}), 400
            company.company_name = requested_company_name
            user.company_name = requested_company_name
        elif company.company_name and normalize_company_name(user.company_name) != normalize_company_name(company.company_name):
            user.company_name = company.company_name

        if 'english_name' in data:
            company.english_name = data['english_name']
        if 'short_name' in data:
            company.short_name = data['short_name']
        if 'credit_code' in data:
            company.credit_code = data['credit_code']
        if 'established_date' in data:
            company.established_date = data['established_date']
        if 'industry' in data:
            company.industry = data['industry']
        if 'website' in data:
            company.website = data['website']
        if 'phone' in data:
            company.phone = data['phone']
        if 'email' in data:
            company.email = data['email']
        if 'address' in data:
            company.address = data['address']
        if 'description' in data:
            company.description = data['description']
        if 'recruitment_manifesto' in data:
            company.recruitment_manifesto = data['recruitment_manifesto']
        elif 'recruitment_declaration' in data:
            company.recruitment_manifesto = data['recruitment_declaration']
        if 'logo_path' in data:
            company.logo_path = data['logo_path']
        if 'enable_employer_brand' in data:
            company.enable_employer_brand = data['enable_employer_brand']
        if 'culture_video_url' in data:
            company.culture_video_url = data['culture_video_url']
        if 'employee_stories' in data:
            company.employee_stories = data['employee_stories']
        if 'resume_retention_days' in data:
            company.resume_retention_days = data['resume_retention_days']
        if 'interview_record_archive_days' in data:
            company.interview_record_archive_days = data['interview_record_archive_days']
        if 'gdpr_compliance' in data:
            company.gdpr_compliance = data['gdpr_compliance']
        if 'consent_popup_enabled' in data:
            company.consent_popup_enabled = data['consent_popup_enabled']
        if 'data_export_enabled' in data:
            company.data_export_enabled = data['data_export_enabled']
        if 'data_deletion_enabled' in data:
            company.data_deletion_enabled = data['data_deletion_enabled']
        if 'audit_log_enabled' in data:
            company.audit_log_enabled = data['audit_log_enabled']
        if 'audit_log_retention_days' in data:
            company.audit_log_retention_days = data['audit_log_retention_days']
        if 'phone_mask_enabled' in data:
            company.phone_mask_enabled = data['phone_mask_enabled']
        if 'email_mask_enabled' in data:
            company.email_mask_enabled = data['email_mask_enabled']
        
        db_session.commit()
        
        # 记录审计日志
        if company.audit_log_enabled:
            audit_log = AuditLog(
                company_id=company.id,
                user_id=user.id,
                action='UPDATE_SETTINGS',
                module='company_settings',
                details=f'更新公司设置：{list(data.keys())}',
                ip_address=request.remote_addr
            )
            db_session.add(audit_log)
            db_session.commit()
        
        company_data = company.to_dict()
        company_data['recruitment_declaration'] = company_data.get('recruitment_manifesto')
        return jsonify({'success': True, 'message': '保存成功', 'company': company_data})


# ==================== 企业管理员相关 API ====================

@app.route('/api/company/dashboard')
def company_dashboard():
    """企业管理员 - 控制台统计"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin or not user.company_name:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        company_name = user.company_name
        
        # 统计数据（单次分组查询，减少重复 count）
        grouped_status = db_session.query(
            JobApplication.status,
            func.count(JobApplication.id)
        ).filter_by(
            company_name=company_name
        ).group_by(
            JobApplication.status
        ).all()

        status_map = {status: count for status, count in grouped_status}
        total = sum(status_map.values())
        pending = status_map.get('submitted', 0)
        interview = status_map.get('interview', 0)
        offer = status_map.get('offer', 0)
        
        return jsonify({
            'success': True,
            'stats': {
                'total': total,
                'pending': pending,
                'interview': interview,
                'offer': offer
            }
        })


@app.route('/api/company/applications')
def company_applications():
    """企业管理员 - 投递列表"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin or not user.company_name:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        company_name = user.company_name
        
        # 获取该公司的所有投递记录
        applications = db_session.query(JobApplication).filter_by(
            company_name=company_name
        ).order_by(JobApplication.applied_at.desc()).limit(200).all()
        
        resume_ids = {app.resume_id for app in applications if app.resume_id}
        user_ids = {app.user_id for app in applications if app.user_id}

        resume_map = {}
        if resume_ids:
            resumes = db_session.query(Resume).filter(Resume.id.in_(resume_ids)).all()
            resume_map = {resume.id: resume for resume in resumes}

        user_map = {}
        if user_ids:
            applicants = db_session.query(User).filter(User.id.in_(user_ids)).all()
            user_map = {applicant.id: applicant for applicant in applicants}

        result = []
        for app in applications:
            resume = resume_map.get(app.resume_id)
            applicant = user_map.get(app.user_id)

            app_data = {
                'id': app.id,
                'user_id': app.user_id,
                'name': app.applicant_name or (resume.name if resume else (applicant.username if applicant else '未知')),
                'phone': app.applicant_phone or (resume.phone if resume else (applicant.phone if applicant else '-')),
                'email': app.applicant_email or (resume.email if resume else (applicant.email if applicant else '-')),
                'job_id': app.job_id,
                'job_name': app.job_name,
                'company_name': app.company_name,
                'status': app.status,
                'applied_at': app.applied_at.isoformat() if app.applied_at else None,
                'resume_url': '/api/resume/download/' + str(app.resume_id) if resume and resume.has_attachment else None
            }
            result.append(app_data)
        
        return jsonify({
            'success': True,
            'applications': result
        })


@app.route('/api/company/jobs')
def company_jobs():
    """企业管理员 - 在招职位"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin or not user.company_name:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        company_name = user.company_name
        
        # 获取分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))
        
        # 获取该公司的所有职位（分页）
        jobs_query = db_session.query(JobPosition).filter_by(
            company_name=company_name
        ).order_by(JobPosition.id.desc())
        
        total = jobs_query.count()
        jobs = jobs_query.offset((page - 1) * per_page).limit(per_page).all()
        
        job_ids = [job.id for job in jobs]
        application_count_map = {}
        if job_ids:
            count_rows = db_session.query(
                JobApplication.job_id,
                func.count(JobApplication.id)
            ).filter(
                JobApplication.job_id.in_(job_ids)
            ).group_by(
                JobApplication.job_id
            ).all()
            application_count_map = {job_id: count for job_id, count in count_rows}

        result = []
        for job in jobs:
            result.append({
                'id': job.id,
                'job_name': job.job_name,
                'min_salary': job.min_salary,
                'max_salary': job.max_salary,
                'city': job.city,
                'education': job.education,
                'experience': job.experience,
                'status': job.status,
                'application_count': application_count_map.get(job.id, 0)
            })
        
        return jsonify({
            'success': True,
            'jobs': result,
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': (total + per_page - 1) // per_page
        })


@app.route('/api/company/application/<int:app_id>/status', methods=['PUT'])
def company_update_application_status(app_id):
    """企业管理员 - 更新申请状态"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403

        user_company_name = (user.company_name or '').strip()
        if not user_company_name:
            return jsonify({'success': False, 'message': '企业账号未绑定公司，无法更新投递状态'}), 403
        user_company_key = normalize_company_name(user_company_name)
        
        data = request.get_json(silent=True) or {}
        new_status = (data.get('status', '') or '').strip().lower()
        interview_context = extract_interview_context_from_request(data) if new_status == 'interview' else {}
        
        if new_status not in ['submitted', 'viewed', 'interview', 'offer', 'rejected']:
            return jsonify({'success': False, 'message': '无效的状态'}), 400

        if new_status == 'interview':
            interview_time_text = (interview_context.get('interview_time_text') or '').strip()
            interview_location = (interview_context.get('interview_location') or '').strip()
            interview_contact = (interview_context.get('interview_contact') or '').strip()

            if not interview_time_text or not interview_location or not interview_contact:
                return jsonify({
                    'success': False,
                    'message': '发起面试需一次性填写完整：面试时间、面试地点、联系电话'
                }), 400
        
        application = db_session.query(JobApplication).filter_by(id=app_id).first()
        
        if not application:
            return jsonify({'success': False, 'message': '申请不存在'}), 404

        # 确保只能修改本公司的申请（允许公司名存在空白/大小写差异）。
        app_company_key = normalize_company_name(application.company_name)
        can_modify = bool(app_company_key) and (app_company_key == user_company_key)

        # 兜底：若投递记录公司名异常，按岗位归属判断权限。
        if not can_modify and application.job_id:
            owner_job = db_session.query(JobPosition).filter_by(id=application.job_id).first()
            if owner_job and normalize_company_name(owner_job.company_name) == user_company_key:
                can_modify = True

        if not can_modify:
            return jsonify({'success': False, 'message': '无权限修改'}), 403
        
        old_status = (application.status or '').strip().lower()
        application.status = new_status

        should_notify = old_status != new_status
        if new_status == 'interview':
            # 面试安排可能发生补充或变更：即使状态未变化，也应重新推送最新时间/地点/联系电话。
            should_notify = True
        if not should_notify and new_status == 'rejected':
            # 企业端重复点击“拒绝”时也触发一次通知，避免前端误以为未生效。
            should_notify = True

        notification_payload = None
        if should_notify:
            notification_payload = build_application_notification_payload(
                db_session,
                application,
                interview_context=interview_context
            )

        db_session.commit()

        notification_pushed = False
        if notification_payload:
            notification_pushed = push_application_status_notification(notification_payload, user.id)
            if not notification_pushed:
                app.logger.warning('状态已更新但通知写入失败: app_id=%s', app_id)
        
        return jsonify({
            'success': True,
            'message': '状态已更新',
            'notification_pushed': notification_pushed
        })


@app.route('/api/company/statistics')
def company_statistics():
    """企业管理员 - 详细统计数据"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    from datetime import datetime, timedelta
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin or not user.company_name:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        company_name = user.company_name
        now = datetime.now()
        
        # 获取分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))
        
        # 本周和月初时间
        week_start = now - timedelta(days=now.weekday())
        month_start = now.replace(day=1)
        
        # 获取所有投递
        all_apps = db_session.query(JobApplication).filter_by(
            company_name=company_name
        ).all()
        
        # 统计数据
        total = len(all_apps)
        week_new = 0
        month_new = 0
        status_count = {
            'submitted': 0,
            'viewed': 0,
            'interview': 0,
            'offer': 0,
            'rejected': 0
        }
        apps_by_job = defaultdict(list)
        processed_days_total = 0
        processed_count = 0

        for app in all_apps:
            if app.applied_at and app.applied_at >= week_start:
                week_new += 1
            if app.applied_at and app.applied_at >= month_start:
                month_new += 1

            if app.status in status_count:
                status_count[app.status] += 1

            if app.job_id is not None:
                apps_by_job[app.job_id].append(app)

            # 简化计算，假设处理时间为当前时间减去申请时间
            if app.status != 'submitted' and app.applied_at:
                processed_days_total += (now - app.applied_at).days
                processed_count += 1

        avg_process_days = round(processed_days_total / processed_count) if processed_count > 0 else 0
        
        # 简历完善率
        resume_ids = set(app.resume_id for app in all_apps if app.resume_id)
        complete_resumes = db_session.query(Resume).filter(
            Resume.id.in_(resume_ids),
            Resume.is_complete == 1
        ).count()
        resume_complete_rate = round(complete_resumes / len(resume_ids) * 100) if resume_ids else 0
        
        # 职位统计（分页）
        jobs_query = db_session.query(JobPosition).filter_by(company_name=company_name)
        total_jobs = jobs_query.count()
        jobs = jobs_query.order_by(JobPosition.id.desc()).offset((page - 1) * per_page).limit(per_page).all()
        
        job_stats = []
        for job in jobs:
            job_apps = apps_by_job.get(job.id, [])
            job_total = len(job_apps)
            job_submitted = sum(1 for app in job_apps if app.status == 'submitted')
            job_interview = sum(1 for app in job_apps if app.status == 'interview')
            job_offer = sum(1 for app in job_apps if app.status == 'offer')
            job_rejected = sum(1 for app in job_apps if app.status == 'rejected')
            conversion_rate = round(job_offer / job_total * 100) if job_total > 0 else 0
            
            job_stats.append({
                'job_name': job.job_name,
                'total': job_total,
                'submitted': job_submitted,
                'interview': job_interview,
                'offer': job_offer,
                'rejected': job_rejected,
                'conversion_rate': conversion_rate
            })
        
        return jsonify({
            'success': True,
            'statistics': {
                'week_new': week_new,
                'month_new': month_new,
                'avg_process_days': avg_process_days,
                'resume_complete_rate': resume_complete_rate,
                'status_count': status_count,
                'job_stats': job_stats,
                'total_jobs': total_jobs,
                'page': page,
                'per_page': per_page,
                'total_pages': (total_jobs + per_page - 1) // per_page
            }
        })


# ==================== 管理员相关 API ====================

@app.route('/api/admin/check')
def check_admin():
    """检查是否为管理员"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        return jsonify({'success': True, 'user': user.to_dict()})


@app.route('/api/admin/dashboard')
def admin_dashboard():
    """管理员控制台 - 数据统计"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401

    now = datetime.now()
    cache_payload = ADMIN_DASHBOARD_CACHE.get('payload')
    cache_expires_at = ADMIN_DASHBOARD_CACHE.get('expires_at')
    if cache_payload and cache_expires_at and now < cache_expires_at:
        return jsonify(cache_payload)

    with get_db_session() as db_session:
        # 统计数据（减少查询次数）
        total_users = db_session.query(User).filter(User.is_admin == False).count()
        total_applications = db_session.query(JobApplication).count()

        grouped_jobs = db_session.query(
            JobPosition.status,
            func.count(JobPosition.id)
        ).group_by(
            JobPosition.status
        ).all()
        status_map = {status: count for status, count in grouped_jobs}
        total_jobs = sum(status_map.values())
        pending_jobs = status_map.get('pending', 0)

        payload = {
            'success': True,
            'admin_username': session.get('username', 'admin'),
            'stats': {
                'total_users': total_users,
                'total_jobs': total_jobs,
                'total_applications': total_applications,
                'pending_jobs': pending_jobs
            }
        }

        ADMIN_DASHBOARD_CACHE['payload'] = payload
        ADMIN_DASHBOARD_CACHE['expires_at'] = now + timedelta(seconds=ADMIN_DASHBOARD_CACHE_TTL_SECONDS)

        return jsonify(payload)


@app.route('/api/admin/users')
def admin_users():
    """管理员 - 用户列表"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 7, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))
        keyword = request.args.get('keyword', '')
        
        query = db_session.query(User).filter(User.is_admin == False)
        if keyword:
            query = query.filter(
                (User.username.like(f'%{keyword}%')) |
                (User.phone.like(f'%{keyword}%')) |
                (User.email.like(f'%{keyword}%'))
            )
        
        total = query.count()
        users = query.order_by(User.id.desc()).offset((page - 1) * per_page).limit(per_page).all()
        
        return jsonify({
            'success': True,
            'users': [u.to_dict() for u in users],
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': (total + per_page - 1) // per_page
        })


@app.route('/api/admin/users/<int:user_id>')
def admin_user_detail(user_id):
    """管理员 - 用户详情"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        target_user = db_session.query(User).filter_by(id=user_id).first()
        if not target_user:
            return jsonify({'success': False, 'message': '用户不存在'}), 404
        
        return jsonify({
            'success': True,
            'user': target_user.to_dict()
        })


@app.route('/api/admin/users/<int:user_id>/status', methods=['PUT'])
def admin_update_user_status(user_id):
    """管理员 - 修改用户状态"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        admin = db_session.query(User).filter_by(id=session['user_id']).first()
        if not admin or not admin.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        data = request.get_json(silent=True) or {}
        target_user = db_session.query(User).filter_by(id=user_id).first()
        
        if not target_user:
            return jsonify({'success': False, 'message': '用户不存在'}), 404
        
        target_user.is_active = coerce_bool(data.get('is_active', True))
        db_session.commit()
        
        return jsonify({'success': True, 'message': '操作成功'})


@app.route('/api/admin/jobs')
def admin_jobs():
    """管理员 - 职位列表"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        # 获取分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 30, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))
        status = request.args.get('status', '')
        
        # 查询
        query = db_session.query(JobPosition)
        if status:
            query = query.filter(JobPosition.status == status)
        
        # 总数（限制最多 500 条）
        total = min(query.count(), 500)
        
        # 分页查询（限制 500 条）
        offset = (page - 1) * per_page
        remaining = max(0, 500 - offset)
        page_limit = min(per_page, remaining)
        jobs = query.order_by(JobPosition.id.desc()).limit(page_limit).offset(offset).all() if page_limit > 0 else []
        
        return jsonify({
            'success': True,
            'jobs': [{
                'id': job.id,
                'job_name': job.job_name,
                'company_name': job.company_name,
                'min_salary': job.min_salary,
                'max_salary': job.max_salary,
                'city': job.city,
                'status': job.status,
                'publish_date': job.publish_date
            } for job in jobs],
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': (total + per_page - 1) // per_page
        })


@app.route('/api/admin/jobs/<int:job_id>/audit', methods=['PUT'])
def admin_audit_job(job_id):
    """管理员 - 审核职位"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        admin = db_session.query(User).filter_by(id=session['user_id']).first()
        if not admin or not admin.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        data = request.get_json(silent=True) or {}
        job = db_session.query(JobPosition).filter_by(id=job_id).first()
        
        if not job:
            return jsonify({'success': False, 'message': '职位不存在'}), 404
        
        result = data.get('result', 'approved')
        reason = data.get('reason', '')
        
        if result == 'approved':
            job.status = 'active'
        elif result == 'rejected':
            job.status = 'rejected'
            if reason:
                # 使用 markupsafe 转义 HTML 字符，防止 XSS 攻击
                safe_reason = str(escape(reason))
                job.description = (job.description or '') + f'\n\n【审核驳回原因】{safe_reason}'
        
        db_session.commit()
        invalidate_admin_dashboard_cache()
        
        return jsonify({'success': True, 'message': '审核完成'})


@app.route('/api/admin/applications')
def admin_applications():
    """管理员 - 投递记录"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 5, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))

        query = db_session.query(JobApplication)
        total = query.count()
        applications = query.order_by(JobApplication.id.desc()).offset((page - 1) * per_page).limit(per_page).all()
        user_ids = list({app.user_id for app in applications if app.user_id is not None})
        users = db_session.query(User).filter(User.id.in_(user_ids)).all() if user_ids else []
        user_map = {u.id: u.username for u in users}
        
        return jsonify({
            'success': True,
            'applications': [{
                'id': app.id,
                'user_id': app.user_id,
                'user_name': user_map.get(app.user_id, '未知'),
                'job_id': app.job_id,
                'job_name': app.job_name,
                'company_name': app.company_name,
                'status': app.status,
                'applied_at': app.applied_at.isoformat() if app.applied_at else None
            } for app in applications],
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': (total + per_page - 1) // per_page
        })


@app.route('/api/admin/password', methods=['PUT'])
def admin_update_password():
    """管理员 - 修改密码"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    from werkzeug.security import generate_password_hash
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        data = request.get_json(silent=True) or {}
        password = data.get('password', '')
        
        if not password or len(password) < 8:
            return jsonify({'success': False, 'message': '密码长度至少 8 位'}), 400
        
        user.password_hash = generate_password_hash(password)
        db_session.commit()
        
        return jsonify({'success': True, 'message': '密码修改成功'})


# ==================== 企业管理模块 ====================

@app.route('/api/admin/companies')
def admin_companies():
    """管理员 - 企业列表"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        # 获取分页参数
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 30, type=int)
        page = max(page, 1)
        per_page = max(1, min(per_page, 100))
        
        base_query = db_session.query(
            JobPosition.company_name.label('company_name'),
            func.max(JobPosition.industry).label('industry'),
            func.max(JobPosition.city).label('city'),
            func.count(JobPosition.id).label('job_count')
        ).filter(
            JobPosition.company_name.isnot(None),
            JobPosition.company_name != ''
        ).group_by(
            JobPosition.company_name
        ).order_by(
            JobPosition.company_name.asc()
        )

        # 总数（限制最多 500 条）
        total = min(base_query.count(), 500)

        # 分页查询（限制 500 条）
        offset = (page - 1) * per_page
        remaining = max(0, 500 - offset)
        page_limit = min(per_page, remaining)

        companies_page = base_query.offset(offset).limit(page_limit).all() if page_limit > 0 else []

        company_list = []
        for company_row in companies_page:
            company_list.append({
                'company_name': company_row.company_name,
                'industry': company_row.industry or '未设置',
                'city': company_row.city or '未设置',
                'job_count': int(company_row.job_count or 0)
            })
        
        return jsonify({
            'success': True,
            'companies': company_list,
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': (total + per_page - 1) // per_page
        })


@app.route('/api/admin/companies/<company_name>')
def admin_company_detail(company_name):
    """管理员 - 企业详情"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        jobs = db_session.query(JobPosition).filter_by(
            company_name=company_name
        ).all()
        
        return jsonify({
            'success': True,
            'company': {
                'company_name': company_name,
                'jobs': [{
                    'id': job.id,
                    'job_name': job.job_name,
                    'salary': f'{job.min_salary}-{job.max_salary}',
                    'city': job.city,
                    'status': job.status
                } for job in jobs]
            }
        })


# ==================== 内容安全与风控 ====================

@app.route('/api/admin/sensitive-words', methods=['GET', 'POST'])
def admin_sensitive_words():
    """管理员 - 敏感词管理"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        if request.method == 'POST':
            data = request.get_json(silent=True) or {}
            word = str(data.get('word', '')).strip()
            word_type = str(data.get('type', '违禁词')).strip() or '违禁词'

            if not word:
                return jsonify({'success': False, 'message': '敏感词不能为空'}), 400
            if len(word) > 100:
                return jsonify({'success': False, 'message': '敏感词长度不能超过 100'}), 400

            words = load_sensitive_words()
            if any(item.get('word') == word and item.get('type') == word_type for item in words):
                return jsonify({'success': False, 'message': '该敏感词已存在'}), 409

            next_id = max([item.get('id', 0) for item in words], default=0) + 1
            new_item = {
                'id': next_id,
                'word': word,
                'type': word_type,
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            words.append(new_item)

            if not save_sensitive_words(words):
                return jsonify({'success': False, 'message': '保存失败，请稍后重试'}), 500

            return jsonify({'success': True, 'message': '敏感词已添加', 'word': new_item})

        words = sorted(load_sensitive_words(), key=lambda item: item.get('id', 0), reverse=True)
        return jsonify({'success': True, 'words': words})


@app.route('/api/admin/jobs/audit/batch', methods=['POST'])
def admin_batch_audit():
    """管理员 - 批量审核职位"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        data = request.get_json(silent=True) or {}
        job_ids = data.get('job_ids', [])
        result = data.get('result', 'approved')
        reason = str(data.get('reason', '')).strip()

        if not isinstance(job_ids, list) or not job_ids:
            return jsonify({'success': False, 'message': '请提供需要审核的职位 ID 列表'}), 400

        if result not in ['approved', 'rejected']:
            return jsonify({'success': False, 'message': '审核结果参数无效'}), 400

        normalized_job_ids = []
        for job_id in job_ids:
            try:
                normalized_job_ids.append(int(job_id))
            except (TypeError, ValueError):
                continue

        if not normalized_job_ids:
            return jsonify({'success': False, 'message': '职位 ID 列表无效'}), 400

        unique_job_ids = list(set(normalized_job_ids))
        jobs = db_session.query(JobPosition).filter(JobPosition.id.in_(unique_job_ids)).all()

        for job in jobs:
            job.status = 'active' if result == 'approved' else 'rejected'
            if result == 'rejected' and reason:
                safe_reason = str(escape(reason))
                job.description = (job.description or '') + f'\n\n【审核驳回原因】{safe_reason}'

        updated_count = len(jobs)
        
        db_session.commit()
        invalidate_admin_dashboard_cache()
        return jsonify({'success': True, 'message': f'批量审核完成，成功更新 {updated_count} 个职位'})


# ==================== 数据导出功能 ====================

@app.route('/api/admin/export/users')
def admin_export_users():
    """管理员 - 导出用户数据"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        users = db_session.query(User).all()
        
        # 生成 CSV 数据（手机号脱敏）
        import csv
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['ID', '用户名', '邮箱', '手机（脱敏）', '学历', '毕业年份', '注册时间'])
        for u in users:
            # 手机号脱敏：保留前 3 位和后 4 位
            masked_phone = ''
            if u.phone and len(u.phone) >= 7:
                masked_phone = u.phone[:3] + '****' + u.phone[-4:]
            elif u.phone:
                masked_phone = '****'
            writer.writerow([u.id, u.username, u.email, masked_phone, u.education or '', u.graduation_year or '', u.created_at])
        
        return jsonify({
            'success': True,
            'message': '导出成功',
            'data': output.getvalue()
        })


@app.route('/api/admin/export/jobs')
def admin_export_jobs():
    """管理员 - 导出职位数据"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        jobs = db_session.query(JobPosition).all()
        
        # 生成 CSV 数据
        import csv
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['ID', '职位名称', '公司', '薪资范围', '城市', '状态', '发布时间'])
        for job in jobs:
            writer.writerow([job.id, job.job_name, job.company_name, f'{job.min_salary}-{job.max_salary}', job.city, job.status, job.publish_date])
        
        return jsonify({
            'success': True,
            'message': '导出成功',
            'data': output.getvalue()
        })


@app.route('/api/admin/export/applications')
def admin_export_applications():
    """管理员 - 导出投递记录"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        applications = db_session.query(JobApplication).limit(1000).all()
        
        # 生成 CSV 数据（手机号脱敏）
        import csv
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['ID', '求职者姓名', '手机号（脱敏）', '职位名称', '公司', '状态', '投递时间'])
        fallback_user_ids = {
            app.user_id for app in applications
            if app.user_id and (not app.applicant_name or not app.applicant_phone)
        }
        fallback_user_map = {}
        if fallback_user_ids:
            fallback_users = db_session.query(User).filter(User.id.in_(fallback_user_ids)).all()
            fallback_user_map = {u.id: u for u in fallback_users}

        for app in applications:
            # 优先使用申请快照，避免历史导出受用户资料后续修改影响
            need_user_fallback = not app.applicant_name or not app.applicant_phone
            applicant_user = fallback_user_map.get(app.user_id) if need_user_fallback else None

            applicant_phone = app.applicant_phone or (applicant_user.phone if applicant_user else '')
            masked_phone = ''
            if applicant_phone and len(applicant_phone) >= 7:
                masked_phone = applicant_phone[:3] + '****' + applicant_phone[-4:]
            elif applicant_phone:
                masked_phone = '****'

            applicant_name = app.applicant_name or (applicant_user.username if applicant_user else '未知')
            writer.writerow([app.id, applicant_name, masked_phone, app.job_name, app.company_name, app.status, app.applied_at])
        
        return jsonify({
            'success': True,
            'message': '导出成功',
            'data': output.getvalue()
        })


# ==================== 日志审计系统 ====================

@app.route('/api/admin/logs')
def admin_logs():
    """管理员 - 操作日志"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        # 返回示例日志数据（后续需要创建日志表）
        logs = [
            {'id': 1, 'user': 'admin', 'action': '登录系统', 'ip': '192.168.1.100', 'time': '2026-04-01 10:30:00'},
            {'id': 2, 'user': 'admin', 'action': '审核职位', 'ip': '192.168.1.100', 'time': '2026-04-01 11:20:00'},
            {'id': 3, 'user': 'admin', 'action': '导出用户数据', 'ip': '192.168.1.100', 'time': '2026-04-01 14:15:00'}
        ]
        
        return jsonify({'success': True, 'logs': logs})


# ==================== 系统配置管理 ====================

@app.route('/api/admin/settings', methods=['GET', 'PUT'])
def admin_settings():
    """管理员 - 系统设置"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '未登录'}), 401
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_admin:
            return jsonify({'success': False, 'message': '无权限'}), 403
        
        if request.method == 'GET':
            # 返回系统配置
            return jsonify({
                'success': True,
                'settings': {
                    'site_name': '职引未来',
                    'max_resume_size': 5,
                    'enable_audit': True,
                    'maintenance_mode': False,
                    'allow_register': True,
                    'job_expire_days': 30
                }
            })
        
        elif request.method == 'PUT':
            data = request.get_json(silent=True) or {}
            # 保存配置（暂时只返回成功）
            return jsonify({'success': True, 'message': '设置已保存'})


# ==================== 启动配置 ====================

def init_app():
    """初始化应用"""
    print(" [OK] 服务启动成功！")
    print(f"    访问地址：http://localhost:{config.FLASK_CONFIG['PORT']}")


# ==================== 聊天功能 ====================

# 导入聊天模型
from models_chat import Conversation, Message, get_session, init_db as init_chat_db

# 初始化聊天数据库表
try:
    init_chat_db()
except Exception as e:
    app.logger.warning("聊天数据库初始化警告：%s", e)


def parse_interview_datetime(raw_value):
    """将前端传入的面试时间归一化为 datetime 与可读文本。"""
    if raw_value in (None, ''):
        return None, ''

    if isinstance(raw_value, datetime):
        value = raw_value
        return value, value.strftime('%Y-%m-%d %H:%M')

    value_text = str(raw_value).strip()
    if not value_text:
        return None, ''

    normalized_text = value_text.replace('T', ' ').rstrip('Z').strip()
    format_candidates = [
        '%Y-%m-%d %H:%M',
        '%Y/%m/%d %H:%M',
        '%Y-%m-%d %H:%M:%S',
        '%Y/%m/%d %H:%M:%S',
        '%Y-%m-%d %H:%M:%S.%f',
        '%Y/%m/%d %H:%M:%S.%f'
    ]

    for fmt in format_candidates:
        try:
            parsed = datetime.strptime(normalized_text, fmt)
            return parsed, parsed.strftime('%Y-%m-%d %H:%M')
        except ValueError:
            continue

    try:
        parsed = datetime.fromisoformat(value_text.replace('Z', '+00:00'))
        if parsed.tzinfo is not None:
            parsed = parsed.replace(tzinfo=None)
        return parsed, parsed.strftime('%Y-%m-%d %H:%M')
    except ValueError:
        app.logger.warning('面试时间解析失败，按原文展示: %s', value_text)
        return None, value_text[:60]


def extract_interview_context_from_request(data):
    """从请求体提取面试通知字段。"""
    interview_time, interview_time_text = parse_interview_datetime(data.get('interview_time'))
    interview_location = (data.get('interview_location') or '').strip()[:255]
    interview_contact = (data.get('interview_contact') or '').strip()[:80]

    return {
        'interview_time': interview_time,
        'interview_time_text': interview_time_text,
        'interview_location': interview_location,
        'interview_contact': interview_contact
    }


def build_interview_notes(contact, source_notes=None):
    """合并并规范化面试备注字段。"""
    notes = []
    if source_notes:
        notes.append(str(source_notes).strip())
    if contact:
        notes.append(f'联系电话：{contact}')

    merged = '\n'.join([part for part in notes if part])
    return merged[:1000] if merged else None


def build_application_notification_payload(db_session, application, interview_context=None):
    """构建通知快照，避免跨会话读取 ORM 对象导致通知丢失。"""
    user_id = application.user_id

    if not user_id and application.resume_id:
        related_resume = db_session.query(Resume).filter_by(id=application.resume_id).first()
        if related_resume and related_resume.user_id:
            user_id = related_resume.user_id

    if not user_id and application.applicant_email:
        related_user = db_session.query(User).filter_by(email=application.applicant_email).first()
        if related_user:
            user_id = related_user.id

    if not user_id and application.applicant_phone:
        related_user = db_session.query(User).filter_by(phone=application.applicant_phone).first()
        if related_user:
            user_id = related_user.id

    if not user_id:
        app.logger.warning('状态通知跳过：无法定位候选人 user_id，app_id=%s', application.id)
        return None

    company_name = (application.company_name or '').strip()
    if not company_name:
        app.logger.warning('状态通知跳过：company_name 为空，app_id=%s', application.id)
        return None

    payload = {
        'app_id': application.id,
        'user_id': int(user_id),
        'company_name': company_name,
        'job_id': application.job_id,
        'job_name': (application.job_name or '').strip(),
        'status': (application.status or '').strip().lower(),
        'interview_time': None,
        'interview_time_text': '',
        'interview_location': '',
        'interview_contact': ''
    }

    if payload['status'] == 'interview':
        context = interview_context or {}
        payload['interview_time'] = context.get('interview_time')
        payload['interview_time_text'] = (context.get('interview_time_text') or '').strip()
        payload['interview_location'] = (context.get('interview_location') or '').strip()
        payload['interview_contact'] = (context.get('interview_contact') or '').strip()

    return payload


def build_application_status_notification_content(payload):
    """根据投递状态生成流程通知文案。"""
    status = (payload.get('status') or '').strip().lower()
    job_name = (payload.get('job_name') or '该岗位').strip() or '该岗位'
    company_name = (payload.get('company_name') or '企业').strip() or '企业'

    if status == 'interview':
        detail_lines = []
        interview_time_text = (payload.get('interview_time_text') or '').strip()
        interview_location = (payload.get('interview_location') or '').strip()
        interview_contact = (payload.get('interview_contact') or '').strip()

        if interview_time_text:
            detail_lines.append(f'面试时间：{interview_time_text}')
        if interview_location:
            detail_lines.append(f'面试地点：{interview_location}')
        if interview_contact:
            detail_lines.append(f'联系电话：{interview_contact}')

        detail_suffix = '\n' + '\n'.join(detail_lines) if detail_lines else '\n请留意后续安排。'
        return f"【流程通知】{company_name}：你投递的“{job_name}”已进入面试阶段。{detail_suffix}"

    status_text_map = {
        'viewed': '已被查看。',
        'offer': '已进入录用阶段，恭喜你。',
        'rejected': '未通过本次筛选，感谢你的投递。',
        'submitted': '状态已更新。'
    }

    status_text = status_text_map.get(status)
    if not status_text:
        return None

    return f"【流程通知】{company_name}：你投递的“{job_name}”{status_text}"


def ensure_notification_conversation(chat_session, payload, company_operator_id):
    """为投递状态通知准备会话，不存在则自动创建。"""
    user_id = payload.get('user_id')
    company_name = (payload.get('company_name') or '').strip()
    job_id = payload.get('job_id')
    job_name = payload.get('job_name')

    if not user_id or not company_name:
        return None

    conversation_query = chat_session.query(Conversation).filter(
        Conversation.user_id == user_id,
        Conversation.company_name == company_name
    )

    if job_id:
        conversation_query = conversation_query.filter(Conversation.job_id == job_id)

    conversation = conversation_query.order_by(Conversation.updated_at.desc()).first()

    if conversation:
        if not conversation.is_active:
            conversation.is_active = True
        if not conversation.job_id and job_id:
            conversation.job_id = job_id
        if not conversation.job_name and job_name:
            conversation.job_name = job_name
        if not conversation.company_id and company_operator_id:
            conversation.company_id = company_operator_id
        return conversation

    conversation = Conversation(
        user_id=user_id,
        company_id=company_operator_id or 0,
        company_name=company_name,
        job_id=job_id,
        job_name=job_name,
        is_active=True
    )
    chat_session.add(conversation)
    chat_session.flush()
    return conversation


def push_application_status_notification(payload, company_operator_id):
    """将投递状态变更写入消息中心通知。"""
    notification_content = build_application_status_notification_content(payload)
    if not notification_content:
        return False

    chat_session = get_session()
    try:
        conversation = ensure_notification_conversation(chat_session, payload, company_operator_id)
        if not conversation:
            return False

        now = datetime.now()

        notification_message = Message(
            conversation_id=conversation.id,
            sender_type='company',
            sender_id=company_operator_id or 0,
            message_type='notification',
            content=notification_content,
            interview_job_id=payload.get('job_id'),
            interview_time=payload.get('interview_time'),
            interview_location=(payload.get('interview_location') or '')[:255] or None,
            interview_notes=build_interview_notes(
                payload.get('interview_contact'),
                payload.get('interview_notes')
            ),
            created_at=now
        )

        chat_session.add(notification_message)

        conversation.updated_at = now
        conversation.last_message_at = now
        conversation.last_message_content = notification_content
        conversation.user_unread_count = (conversation.user_unread_count or 0) + 1
        conversation.company_unread_count = 0

        chat_session.commit()
        return True
    except Exception:
        chat_session.rollback()
        app.logger.exception("写入状态通知失败，尝试文本兜底: app_id=%s", payload.get('app_id'))
        return push_application_status_notification_as_text(
            payload,
            company_operator_id,
            notification_content
        )
    finally:
        chat_session.close()


def push_application_status_notification_as_text(payload, company_operator_id, notification_content):
    """通知消息写入失败时，退化为普通文本消息，保证用户可见。"""
    chat_session = get_session()
    try:
        conversation = ensure_notification_conversation(chat_session, payload, company_operator_id)
        if not conversation:
            return False

        now = datetime.now()
        fallback_content = notification_content.replace('【流程通知】', '[流程通知] ')

        fallback_message = Message(
            conversation_id=conversation.id,
            sender_type='company',
            sender_id=company_operator_id or 0,
            message_type='text',
            content=fallback_content,
            interview_job_id=payload.get('job_id'),
            interview_time=payload.get('interview_time'),
            interview_location=(payload.get('interview_location') or '')[:255] or None,
            interview_notes=build_interview_notes(
                payload.get('interview_contact'),
                payload.get('interview_notes')
            ),
            created_at=now
        )

        chat_session.add(fallback_message)

        conversation.updated_at = now
        conversation.last_message_at = now
        conversation.last_message_content = fallback_content
        conversation.user_unread_count = (conversation.user_unread_count or 0) + 1
        conversation.company_unread_count = 0

        chat_session.commit()
        return True
    except Exception:
        chat_session.rollback()
        app.logger.exception("文本兜底通知也写入失败: app_id=%s", payload.get('app_id'))
        return False
    finally:
        chat_session.close()


@app.route('/api/chat/conversations', methods=['GET'])
def get_conversations():
    """获取会话列表"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))
            
            if is_company_mode:
                if not user.is_company_admin:
                    return jsonify({'success': False, 'message': '当前不是企业登录状态'}), 403
                conversations = chat_session.query(Conversation).filter(
                    Conversation.company_name == user.company_name,
                    Conversation.is_active == True
                ).order_by(Conversation.updated_at.desc()).all()
            else:
                conversations = chat_session.query(Conversation).filter(
                    Conversation.user_id == session['user_id'],
                    Conversation.is_active == True
                ).order_by(Conversation.updated_at.desc()).all()

            if not conversations:
                return jsonify({'success': True, 'conversations': []})

            conversation_user_ids = {conv.user_id for conv in conversations if conv.user_id}
            conversation_company_names = {conv.company_name for conv in conversations if conv.company_name}
            conversation_ids = [conv.id for conv in conversations]

            related_applications = []
            if conversation_user_ids and conversation_company_names:
                related_applications = db_session.query(JobApplication).filter(
                    JobApplication.user_id.in_(conversation_user_ids),
                    JobApplication.company_name.in_(conversation_company_names)
                ).order_by(JobApplication.applied_at.asc()).all()

            latest_application_by_key = {}
            latest_application_by_company = {}
            earliest_application_by_company = {}
            for application in related_applications:
                key = (application.user_id, application.company_name, application.job_id)
                company_key = (application.user_id, application.company_name)

                latest_application_by_key[key] = application
                latest_application_by_company[company_key] = application
                earliest_application_by_company.setdefault(company_key, application)

            resume_ids = {app.resume_id for app in related_applications if app.resume_id}
            resume_name_map = {}
            if resume_ids:
                resumes = db_session.query(Resume).filter(Resume.id.in_(resume_ids)).all()
                resume_name_map = {
                    resume.id: resume.name
                    for resume in resumes
                    if resume and resume.name
                }

            applicant_name_map = {}
            if conversation_user_ids:
                applicants = db_session.query(User).filter(User.id.in_(conversation_user_ids)).all()
                applicant_name_map = {
                    applicant.id: applicant.username
                    for applicant in applicants
                    if applicant and applicant.username
                }

            latest_message_map = {}
            if conversation_ids:
                latest_message_ids = [
                    row.max_message_id
                    for row in chat_session.query(
                        func.max(Message.id).label('max_message_id')
                    ).filter(
                        Message.conversation_id.in_(conversation_ids)
                    ).group_by(
                        Message.conversation_id
                    ).all()
                    if row.max_message_id
                ]

                if latest_message_ids:
                    latest_messages = chat_session.query(Message).filter(
                        Message.id.in_(latest_message_ids)
                    ).all()
                    latest_message_map = {
                        message.conversation_id: message
                        for message in latest_messages
                    }

            result = []
            for conv in conversations:
                conv_data = conv.to_dict()

                if conv.job_id:
                    matched_application = latest_application_by_key.get((conv.user_id, conv.company_name, conv.job_id))
                else:
                    # 无职位绑定时，取最早申请以保持会话展示名称稳定
                    matched_application = earliest_application_by_company.get((conv.user_id, conv.company_name))

                if not matched_application:
                    matched_application = latest_application_by_company.get((conv.user_id, conv.company_name))

                applicant_name = None
                if matched_application and matched_application.applicant_name:
                    applicant_name = matched_application.applicant_name

                if not applicant_name and matched_application and matched_application.resume_id:
                    applicant_name = resume_name_map.get(matched_application.resume_id)

                if not applicant_name:
                    applicant_name = applicant_name_map.get(conv.user_id, f"用户{conv.user_id}")

                conv_data['user_name'] = applicant_name

                # 兼容旧数据：若会话表未维护最后一条消息，则从消息表兜底读取
                if not conv_data.get('last_message_content') or not conv_data.get('last_message_at'):
                    last_message = latest_message_map.get(conv.id)
                    if last_message:
                        if not conv_data.get('last_message_content'):
                            conv_data['last_message_content'] = last_message.content
                        if not conv_data.get('last_message_at'):
                            conv_data['last_message_at'] = (
                                last_message.created_at.isoformat() if last_message.created_at else conv_data.get('updated_at')
                            )

                if not conv_data.get('last_message_at'):
                    conv_data['last_message_at'] = conv_data.get('updated_at')

                result.append(conv_data)

            return jsonify({'success': True, 'conversations': result})
    
    except Exception:
        app.logger.exception("获取会话列表失败")
        return jsonify({'success': False, 'message': '获取失败'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/conversations', methods=['POST'])
def create_conversation():
    """创建新会话"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    data = request.get_json(silent=True) or {}
    company_name = (data.get('company_name') or '').strip()
    job_id = data.get('job_id')
    job_name = (data.get('job_name') or '').strip()
    target_user_id = data.get('user_id')

    try:
        job_id = int(job_id) if job_id not in (None, '', 'null') else None
    except (TypeError, ValueError):
        return jsonify({'success': False, 'message': '职位参数无效'}), 400

    if company_name and len(company_name) > MAX_COMPANY_NAME_LENGTH:
        return jsonify({'success': False, 'message': '公司名称长度超出限制'}), 400

    if job_name and len(job_name) > MAX_JOB_NAME_LENGTH:
        return jsonify({'success': False, 'message': '职位名称长度超出限制'}), 400

    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))

            if is_company_mode:
                if not user.is_company_admin:
                    return jsonify({'success': False, 'message': '当前不是企业登录状态'}), 403
                try:
                    chat_user_id = int(target_user_id) if target_user_id else session['user_id']
                    if chat_user_id <= 0:
                        raise ValueError
                except (TypeError, ValueError):
                    return jsonify({'success': False, 'message': '用户参数无效'}), 400
                bound_company_name = (user.company_name or '').strip()
                if not bound_company_name:
                    return jsonify({'success': False, 'message': '企业账号未绑定公司'}), 403

                # 企业模式下禁止使用请求体覆盖公司名，避免产生跨公司脏会话。
                if company_name and normalize_company_name(company_name) != normalize_company_name(bound_company_name):
                    return jsonify({'success': False, 'message': '无权为其他公司创建会话'}), 403

                company_name = bound_company_name
                company_owner_id = user.id
            else:
                chat_user_id = session['user_id']

                # 普通用户兜底：若前端未传 company_name，尝试从职位表补齐
                if not company_name and job_id:
                    job = db_session.query(JobPosition).filter_by(id=job_id).first()
                    if job:
                        company_name = (job.company_name or '').strip()
                        if not job_name:
                            job_name = job.job_name or ''

                company_owner = None
                if company_name:
                    company_owner = db_session.query(User).filter_by(
                        is_company_admin=True,
                        company_name=company_name
                    ).order_by(User.id.asc()).first()

                if company_owner:
                    company_owner_id = company_owner.id
                else:
                    company_record = db_session.query(Company).filter_by(company_name=company_name).first() if company_name else None
                    company_owner_id = company_record.id if company_record else 0

            if not company_name:
                return jsonify({'success': False, 'message': '公司名称不能为空'}), 400

            # 企业模式：同候选人+同公司+同岗位复用会话；普通用户模式保持同公司复用
            existing_query = chat_session.query(Conversation).filter_by(
                user_id=chat_user_id,
                company_name=company_name
            )
            if is_company_mode and job_id:
                existing_query = existing_query.filter_by(job_id=job_id)

            existing = existing_query.order_by(Conversation.updated_at.desc()).first()
            if existing:
                if not existing.is_active:
                    existing.is_active = True
                # 保留创建时职位信息，但允许在为空时补齐
                if not existing.job_id and job_id:
                    existing.job_id = job_id
                if not existing.job_name and job_name:
                    existing.job_name = job_name
                existing.updated_at = datetime.now()
                chat_session.commit()
                return jsonify({'success': True, 'conversation': existing.to_dict(), 'created': False})

            # 创建新会话
            conversation = Conversation(
                user_id=chat_user_id,
                company_id=company_owner_id,
                company_name=company_name,
                job_id=job_id,
                job_name=job_name
            )
            chat_session.add(conversation)
            chat_session.commit()
            return jsonify({'success': True, 'conversation': conversation.to_dict(), 'created': True})
    except Exception:
        chat_session.rollback()
        app.logger.exception("创建会话失败")
        return jsonify({'success': False, 'message': '创建失败'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/messages/<int:conversation_id>', methods=['GET'])
def get_messages(conversation_id):
    """获取消息列表（支持 since 增量查询 + 自动标记已读）"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    since_id = request.args.get('since', type=int)

    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))

            conversation = chat_session.query(Conversation).filter_by(id=conversation_id).first()

            if not conversation or not conversation.is_active:
                return jsonify({'success': False, 'message': '会话不存在'}), 404

            if is_company_mode:
                if not user.is_company_admin:
                    return jsonify({'success': False, 'message': '当前不是企业登录状态'}), 403
                if conversation.company_name != user.company_name:
                    return jsonify({'success': False, 'message': '无权访问'}), 403
                if (conversation.company_unread_count or 0) > 0:
                    conversation.company_unread_count = 0
                    chat_session.commit()
            else:
                if conversation.user_id != session['user_id']:
                    return jsonify({'success': False, 'message': '无权访问'}), 403
                if (conversation.user_unread_count or 0) > 0:
                    conversation.user_unread_count = 0
                    chat_session.commit()

            conversation_data = conversation.to_dict()

        query = chat_session.query(Message).filter_by(conversation_id=conversation_id)
        if since_id:
            messages = query.filter(Message.id > since_id).order_by(Message.id.asc()).all()
        else:
            # 首次加载只取最近 200 条，再倒序回正，既保证性能也保持时间顺序
            messages = query.order_by(Message.created_at.desc(), Message.id.desc()).limit(200).all()
            messages.reverse()

        # 自动标记对方发送的未读消息为已读
        now = datetime.now()
        unread_messages = [msg for msg in messages if not msg.is_read and msg.sender_type != ('company' if is_company_mode else 'user')]
        if unread_messages:
            for msg in unread_messages:
                msg.is_read = True
                msg.read_at = now
            chat_session.commit()

        result = [msg.to_dict() for msg in messages]
        return jsonify({'success': True, 'messages': result, 'conversation': conversation_data})

    except Exception:
        app.logger.exception("获取消息失败")
        return jsonify({'success': False, 'message': '获取失败'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/conversations/<int:conversation_id>', methods=['GET'])
def get_conversation_detail(conversation_id):
    """兼容企业端旧接口：返回会话消息列表。"""
    return get_messages(conversation_id)


@app.route('/api/chat/messages/<int:message_id>/confirm-interview', methods=['POST'])
def confirm_interview_notification(message_id):
    """候选人在流程通知中确认面试。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            if session.get('is_company') or session.get('is_admin') or user.is_company_admin:
                return jsonify({'success': False, 'message': '仅求职者可确认面试'}), 403

        target_message = chat_session.query(Message).filter_by(id=message_id).first()
        if not target_message:
            return jsonify({'success': False, 'message': '通知不存在'}), 404

        conversation = chat_session.query(Conversation).filter_by(id=target_message.conversation_id).first()
        if not conversation or not conversation.is_active:
            return jsonify({'success': False, 'message': '会话不存在'}), 404

        if conversation.user_id != session['user_id']:
            return jsonify({'success': False, 'message': '无权操作该通知'}), 403

        message_type = (target_message.message_type or '').strip().lower()
        content_text = (target_message.content or '').strip()
        has_interview_hint = bool(
            target_message.interview_time
            or (target_message.interview_location or '').strip()
            or ('面试' in content_text)
        )

        if message_type != 'notification' and '流程通知' not in content_text:
            return jsonify({'success': False, 'message': '当前消息不是流程通知'}), 400

        if not has_interview_hint:
            return jsonify({'success': False, 'message': '该通知不是面试通知'}), 400

        existing_notes = (target_message.interview_notes or '').strip()
        if INTERVIEW_CONFIRMATION_MARKER in existing_notes:
            return jsonify({'success': True, 'message': '你已确认过该面试'})

        now = datetime.now()
        confirmation_text = '我已确认参加面试，感谢安排。'

        if target_message.interview_time:
            confirmation_text += f"（时间：{target_message.interview_time.strftime('%Y-%m-%d %H:%M')}）"
        if target_message.interview_location:
            confirmation_text += f"（地点：{target_message.interview_location}）"

        confirmation_note = f"{INTERVIEW_CONFIRMATION_MARKER}:{now.isoformat()}\n候选人确认时间：{now.strftime('%Y-%m-%d %H:%M')}"
        target_message.interview_notes = '\n'.join(
            [segment for segment in [existing_notes, confirmation_note] if segment]
        )[:1000]
        target_message.updated_at = now

        confirm_message = Message(
            conversation_id=conversation.id,
            sender_type='user',
            sender_id=session['user_id'],
            message_type='text',
            content=confirmation_text,
            interview_job_id=target_message.interview_job_id,
            interview_time=target_message.interview_time,
            interview_location=target_message.interview_location,
            created_at=now
        )
        chat_session.add(confirm_message)

        conversation.updated_at = now
        conversation.last_message_at = now
        conversation.last_message_content = confirmation_text
        conversation.company_unread_count = (conversation.company_unread_count or 0) + 1
        conversation.user_unread_count = 0

        chat_session.commit()
        return jsonify({'success': True, 'message': '面试确认已发送'})
    except Exception:
        chat_session.rollback()
        app.logger.exception('确认面试失败: message_id=%s', message_id)
        return jsonify({'success': False, 'message': '确认失败，请稍后重试'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/conversations/<int:conversation_id>', methods=['DELETE'])
def delete_conversation(conversation_id):
    """删除会话（软删除）。"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            conversation = chat_session.query(Conversation).filter_by(id=conversation_id).first()
            if not conversation or not conversation.is_active:
                return jsonify({'success': False, 'message': '会话不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))
            if is_company_mode:
                if not user.is_company_admin:
                    return jsonify({'success': False, 'message': '当前不是企业登录状态'}), 403
                if conversation.company_name != user.company_name:
                    return jsonify({'success': False, 'message': '无权删除该会话'}), 403
            else:
                if conversation.user_id != session['user_id']:
                    return jsonify({'success': False, 'message': '无权删除该会话'}), 403

            conversation.is_active = False
            conversation.user_unread_count = 0
            conversation.company_unread_count = 0
            conversation.updated_at = datetime.now()
            chat_session.commit()

            return jsonify({'success': True, 'message': '会话已删除'})
    except Exception:
        chat_session.rollback()
        app.logger.exception("删除会话失败")
        return jsonify({'success': False, 'message': '删除失败'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/messages', methods=['POST'])
def send_message():
    """发送消息"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401
    
    data = request.get_json(silent=True) or {}
    conversation_id = data.get('conversation_id')
    content = data.get('content', '').strip()
    message_type = (data.get('message_type') or 'text').strip().lower()
    file_url = data.get('file_url', '').strip()
    duration = data.get('duration', 0)

    if not conversation_id:
        return jsonify({'success': False, 'message': '参数不完整'}), 400

    if message_type == 'voice':
        if not file_url:
            return jsonify({'success': False, 'message': '语音消息缺少文件'}), 400
        if not content:
            content = f'[语音消息] {duration}秒'
    elif not content:
        return jsonify({'success': False, 'message': '参数不完整'}), 400

    if len(content) > MAX_CHAT_MESSAGE_LENGTH:
        return jsonify({'success': False, 'message': '消息内容过长'}), 400

    if message_type not in ALLOWED_CHAT_MESSAGE_TYPES:
        return jsonify({'success': False, 'message': '消息类型无效'}), 400

    try:
        conversation_id = int(conversation_id)
        if conversation_id <= 0:
            raise ValueError
    except (TypeError, ValueError):
        return jsonify({'success': False, 'message': '会话参数无效'}), 400
    
    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))

            conversation = chat_session.query(Conversation).filter_by(id=conversation_id).first()
            
            if not conversation or not conversation.is_active:
                return jsonify({'success': False, 'message': '会话不存在'}), 404
            
            if is_company_mode:
                if not user.is_company_admin:
                    return jsonify({'success': False, 'message': '当前不是企业登录状态'}), 403
                if conversation.company_name != user.company_name:
                    return jsonify({'success': False, 'message': '无权发送消息'}), 403
                sender_type = 'company'
            else:
                if conversation.user_id != session['user_id']:
                    return jsonify({'success': False, 'message': '无权发送消息'}), 403
                sender_type = 'user'
            
            message = Message(
                conversation_id=conversation_id,
                sender_type=sender_type,
                sender_id=session['user_id'],
                message_type=message_type or 'text',
                content=content,
                file_url=file_url if message_type == 'voice' else None,
                file_name=f'voice_{duration}s' if message_type == 'voice' else None
            )
            
            chat_session.add(message)
            
            now = datetime.now()
            conversation.updated_at = now
            conversation.last_message_at = now
            conversation.last_message_content = content

            if sender_type == 'company':
                conversation.user_unread_count = (conversation.user_unread_count or 0) + 1
                conversation.company_unread_count = 0
            else:
                conversation.company_unread_count = (conversation.company_unread_count or 0) + 1
                conversation.user_unread_count = 0
            
            chat_session.commit()
            
            return jsonify({'success': True, 'message': message.to_dict()})
    
    except Exception:
        chat_session.rollback()
        app.logger.exception("发送消息失败")
        return jsonify({'success': False, 'message': '发送失败'}), 500
    finally:
        chat_session.close()


@app.route('/api/chat/voice/upload', methods=['POST'])
def upload_voice_message():
    """上传语音消息文件"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)

    if file_size > 5 * 1024 * 1024:
        return jsonify({'success': False, 'message': '语音文件大小不能超过 5MB'}), 400

    allowed_extensions = {'webm', 'ogg', 'mp3', 'wav', 'm4a'}
    original_filename = file.filename
    if '.' not in original_filename:
        return jsonify({'success': False, 'message': '无效的文件名'}), 400

    ext = original_filename.rsplit('.', 1)[1].lower()
    if ext not in allowed_extensions:
        return jsonify({'success': False, 'message': '仅支持 webm/ogg/mp3/wav/m4a 格式'}), 400

    import uuid
    basedir = os.path.abspath(os.path.dirname(__file__))
    upload_dir = os.path.join(basedir, 'uploads', 'voice')
    os.makedirs(upload_dir, exist_ok=True)

    unique_filename = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(upload_dir, unique_filename)
    file.save(filepath)

    voice_url = f"/api/chat/voice/{unique_filename}"
    duration = request.form.get('duration', type=int, default=0)

    return jsonify({
        'success': True,
        'voice_url': voice_url,
        'duration': duration,
        'file_size': file_size
    })


@app.route('/api/chat/voice/<filename>')
def serve_voice_file(filename):
    """提供语音文件访问"""
    basedir = os.path.abspath(os.path.dirname(__file__))
    voice_dir = os.path.join(basedir, 'uploads', 'voice')
    safe_name = secure_filename(filename)
    filepath = os.path.join(voice_dir, safe_name)
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': '文件不存在'}), 404

    mime_map = {
        'webm': 'audio/webm',
        'ogg': 'audio/ogg',
        'mp3': 'audio/mpeg',
        'wav': 'audio/wav',
        'm4a': 'audio/mp4'
    }
    ext = safe_name.rsplit('.', 1)[-1].lower() if '.' in safe_name else ''
    mime_type = mime_map.get(ext, 'audio/mpeg')

    return send_file(filepath, mimetype=mime_type)


@app.route('/api/chat/messages/search', methods=['GET'])
def search_messages():
    """搜索历史消息"""
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': '请先登录'}), 401

    keyword = request.args.get('keyword', '').strip()
    conversation_id = request.args.get('conversation_id', type=int)
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)

    if not keyword:
        return jsonify({'success': False, 'message': '请输入搜索关键词'}), 400

    if len(keyword) > 100:
        return jsonify({'success': False, 'message': '搜索关键词过长'}), 400

    per_page = min(per_page, 50)
    page = max(page, 1)

    chat_session = get_session()
    try:
        with get_db_session() as db_session:
            user = db_session.query(User).filter_by(id=session['user_id']).first()
            if not user or not user.is_active:
                return jsonify({'success': False, 'message': '用户不存在'}), 404

            is_company_mode = bool(session.get('is_company', False))

        query = chat_session.query(Message).filter(
            Message.content.like(f'%{keyword}%'),
            Message.message_type.in_(['text', 'emoji'])
        )

        if conversation_id:
            conv = chat_session.query(Conversation).filter_by(id=conversation_id).first()
            if not conv or not conv.is_active:
                return jsonify({'success': False, 'message': '会话不存在'}), 404

            if is_company_mode:
                if conv.company_name != user.company_name:
                    return jsonify({'success': False, 'message': '无权访问'}), 403
            else:
                if conv.user_id != session['user_id']:
                    return jsonify({'success': False, 'message': '无权访问'}), 403

            query = query.filter(Message.conversation_id == conversation_id)
        else:
            if is_company_mode:
                conv_ids = [c.id for c in chat_session.query(Conversation).filter_by(
                    company_name=user.company_name, is_active=True
                ).all()]
            else:
                conv_ids = [c.id for c in chat_session.query(Conversation).filter_by(
                    user_id=session['user_id'], is_active=True
                ).all()]

            if not conv_ids:
                return jsonify({'success': True, 'messages': [], 'total': 0, 'page': page})

            query = query.filter(Message.conversation_id.in_(conv_ids))

        total = query.count()
        messages = query.order_by(Message.created_at.desc()).offset(
            (page - 1) * per_page
        ).limit(per_page).all()

        result = []
        for msg in messages:
            msg_dict = msg.to_dict()
            conv = chat_session.query(Conversation).filter_by(id=msg.conversation_id).first()
            msg_dict['conversation_title'] = (
                f"{conv.company_name} · {conv.job_name}" if conv else '未知会话'
            )
            result.append(msg_dict)

        return jsonify({
            'success': True,
            'messages': result,
            'total': total,
            'page': page,
            'per_page': per_page
        })

    except Exception:
        app.logger.exception("搜索消息失败")
        return jsonify({'success': False, 'message': '搜索失败'}), 500
    finally:
        chat_session.close()


@app.route('/chat/company')
def company_chat():
    """企业 HR 聊天页面"""
    if 'user_id' not in session:
        return redirect('/login')

    if not session.get('is_company', False):
        return redirect('/login?message=请使用企业登录方式登录')
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_company_admin:
            return redirect('/login')
        
        return render_template('company_chat.html', user=user)


@app.route('/chat/user')
def user_chat():
    """普通用户聊天页面"""
    if 'user_id' not in session:
        return redirect('/login')

    # 使用本次登录模式判断门户，避免跨角色登录后误跳转登录页
    if session.get('is_admin'):
        return redirect('/admin')
    if session.get('is_company'):
        return redirect('/company')
    
    with get_db_session() as db_session:
        user = db_session.query(User).filter_by(id=session['user_id']).first()
        if not user or not user.is_active:
            session.clear()
            return redirect('/login?message=登录已失效，请重新登录')
        
        return render_template('user_chat.html', user=user)


@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    """AI 职业助手聊天接口。"""
    data = request.get_json(silent=True)
    if not data or not data.get('message'):
        return jsonify({'success': False, 'message': '请输入消息内容'}), 400

    user_msg = data['message'].strip()
    if not user_msg:
        return jsonify({'success': False, 'message': '消息不能为空'}), 400

    ai_config = getattr(config, 'AI_CONFIG', {})
    api_key = str(ai_config.get('API_KEY', '')).strip()
    base_url = str(ai_config.get('BASE_URL', '')).strip().rstrip('/')
    model = str(ai_config.get('MODEL', 'qwen-plus')).strip()

    if not api_key:
        return jsonify({
            'success': False,
            'message': '未配置 AI_API_KEY，请先在 .env 中添加接口密钥'
        }), 503

    if not base_url:
        return jsonify({
            'success': False,
            'message': '未配置 AI_API_BASE_URL，请先在 .env 中添加接口地址'
        }), 503

    if 'dashscope' in base_url.lower() and looks_like_realtime_model(model):
        return jsonify({
            'success': False,
            'message': '当前模型需要 DashScope realtime WebSocket 通道，且主要面向音频/图像输入，不适合当前文本框直接调用。',
            'provider_error': 'model_requires_realtime_transport',
            'transport': 'websocket_realtime',
            'supported_input': ['audio', 'image'],
            'suggestion': '如果要继续使用这个模型，需要增加语音/图像输入入口；如果只想保留文本聊天，请切换到支持 HTTP 对话的模型。',
            'attempts': []
        }), 503

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }

    errors = []
    for candidate in build_ai_request_candidates(base_url, model, user_msg, ai_config):
        try:
            response = requests.post(
                candidate['url'],
                headers=headers,
                json=candidate['payload'],
                timeout=ai_config.get('TIMEOUT', 60)
            )
            response.raise_for_status()
            result = response.json()

            reply = extract_ai_reply(result)
            if reply:
                return jsonify({
                    'success': True,
                    'reply': reply,
                    'is_simulated': False,
                    'model': model,
                    'provider_url': candidate['url']
                })

            errors.append(f"{candidate['url']}: 返回成功但未解析到回复内容")
            app.logger.warning('AI 接口返回成功但未解析到回复内容: %s', result)
        except requests.HTTPError as http_exc:
            detail = extract_provider_error(getattr(http_exc, 'response', None))
            errors.append(f"{candidate['url']}: {http_exc}{' - ' + detail if detail else ''}")
            app.logger.warning('AI 接口候选请求失败: %s', errors[-1])
            if detail and 'does not support http call' in detail.lower():
                continue
        except requests.RequestException as exc:
            errors.append(f"{candidate['url']}: {exc}")
            app.logger.warning('AI 接口候选请求异常: %s', errors[-1])
        except ValueError:
            errors.append(f"{candidate['url']}: 接口返回了非 JSON 内容")
            app.logger.warning('AI 接口候选请求返回非 JSON 内容: %s', candidate['url'])

    app.logger.error('AI 接口所有候选调用均失败: %s', errors)
    return jsonify({
        'success': False,
        'message': 'AI 接口调用失败，请检查模型、权限或接口地址设置',
        'provider_error': errors[-1] if errors else None,
        'attempts': errors
    }), 502


if __name__ == '__main__':
    init_app()
    app.run(
        host=config.FLASK_CONFIG['HOST'],
        port=config.FLASK_CONFIG['PORT'],
        debug=config.FLASK_CONFIG['DEBUG']
    )
