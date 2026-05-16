import os
import re

readme_path = r'c:\Users\床垫子\Desktop\积分单元\README.md'
with open(readme_path, 'r', encoding='utf-8') as f:
    readme_text = f.read()

diagrams = {
    '### 1.5 初始顶层用例模型': '### 1.5 初始顶层用例模型\n\n```mermaid\ngraph LR\n    subgraph 用户角色\n        S[求职者 Student]\n        C[企业人事 HR]\n        A[系统管理员 Admin]\n    end\n    subgraph 职引未来 - 高校毕业生就业服务平台\n        S -->|检索职位| UC1(职位搜索与筛选)\n        S -->|查看分析| UC2(市场数据洞察查看)\n        S -->|上传云端| UC3(简历云端管理解析)\n        S -->|申请登记| UC4(投递追踪)\n        S -->|在线沟通| UC5(即时通讯联络)\n        \n        C -->|发布及维保| UC6(职位发布管理)\n        C -->|审核申请单| UC7(简历申请表审核)\n        C -->|发送消息| UC8(发起面谈沟通)\n        \n        A -->|资质审核| UC9(企业资质全局管控)\n        A -->|封禁违规| UC10(涉敏账号治理)\n        A -->|维保监视| UC11(系统监控审核)\n    end\n```\n\n',

    '#### 系统架构设计': '#### 系统架构设计\n\n```mermaid\ngraph TD\n    subgraph 视图表现层 Views / Frontend\n        UI[Jinja2 & Bootstrap UI]\n        JS[DOM原生层 Fetch API]\n        ECH[ECharts & Pyecharts可视化组件]\n    end\n    subgraph 业务路由层 Controllers\n        AUTH[鉴权及身份横向拦截网 @require_role]\n        API[Restful核心API]\n        VIS[独立外接静态计算缓存引擎 visualizer.py]\n    end\n    subgraph 持久模型层 Models / Data\n        ORM[SQLAlchemy 实体映射总线]\n        DB[(MySQL 8.0 主库)]\n        DATA[Pandas 独立重构任务件 data_cleaner.py]\n    end\n    UI --> AUTH\n    JS --> API\n    ECH --> VIS\n    AUTH --> ORM \n    API --> ORM\n    VIS --> DATA\n    DATA --> DB\n    ORM --> DB\n```\n\n',

    '#### 系统核心投递与沟通流程图': '#### 系统核心投递与沟通流程图\n\n```mermaid\nflowchart TD\n    A([求职者浏览岗位]) --> B{是否已登录并完善简历?}\n    B -- 否 --> C[拦断/跳转至简历中心]\n    C --> B\n    B -- 是 --> D[发起职位投递请求]\n    D --> E[系统校验岗限流/防重发]\n    E --> F[产生 Application 快照]\n    F --> G([HR 收到提醒])\n    G --> H{执行简历一筛评定}\n    H -- 淘汰 --> I[单录改 Refused 归档]\n    H -- 过审 --> J[单录标定 Invited邀约面谈]\n    J --> K[平台解除双方私聊屏蔽墙]\n    K --> L([正式进入即时通讯环节])\n```\n\n',

    '#### 业务领域静态模型（核心实体类图）': '#### 业务领域静态模型（核心实体类图）\n\n```mermaid\nclassDiagram\n    class User {\n        +int id\n        +string email\n        +string password_hash\n        +string role\n        +verify_password()\n    }\n    class StudentProfile {\n        +int user_id\n        +string real_name\n        +string university\n    }\n    class Resume {\n        +int id\n        +string file_path\n    }\n    class Company {\n        +int id\n        +string company_name\n    }\n    class Job {\n        +int id\n        +string title\n        +int salary_min\n    }\n    User "1" *-- "1" StudentProfile : 关联附加\n    StudentProfile "1" *-- "1..*" Resume : 拥有上传件\n    User "1" *-- "0..*" Company : 法权分配\n    Company "1" *-- "1..*" Job : 分发管理挂载\n```\n\n',
    
    '#### 业务领域动态模型（简历投递与邀约状态流转时序图）': '#### 业务领域动态模型（简历投递与邀约状态流转时序图）\n\n```mermaid\nsequenceDiagram\n    participant S as 求职前端\n    participant B as Flask 后端\n    participant DB as MySQL 数据库\n    participant C as HR 后台\n\n    S->>B: POST /apply 发送请求(简历ID,职位ID)\n    B->>DB: 事务挂载及重放请求拦截\n    DB-->>B: 回传验证许可确认\n    B->>DB: 执行写入 Application 表\n    B-->>S: 告知 200 OK 投递完毕\n\n    C->>B: GET /applications 请求待办视图\n    B->>DB: 多表联合调阅相关连信源快照\n    DB-->>B: 传回聚合实体流\n    B-->>C: 推送 JSON 映射表格显示项\n    \n    C->>B: PUT 邀约状态变更为(Invited)\n    B->>DB: 改写状态、激活解扣私聊防线\n    B-->>C: 前端弹窗开启进入 Chat\n```\n\n',

    '#### 数据库 ER 模型图': '#### 数据库 ER 模型图\n\n```mermaid\nerDiagram\n    USER ||--o| STUDENT_PROFILE : "对应1附加配置"\n    USER ||--o| COMPANY : "挂靠一管理主体"\n    COMPANY ||--o{ JOB : "挂职多个活跃单"\n    STUDENT_PROFILE ||--o{ RESUME : "解析出N附件"\n    JOB ||--o{ APPLICATION : "吸收容纳各方来投"\n    RESUME ||--o{ APPLICATION : "支持多元派驻散发"\n    \n    USER { int id PK } \n    JOB { int id PK } \n    APPLICATION { int id PK } \n```\n\n'
}

for title, merm_code in diagrams.items():
    # Remove the previously injected markdown image links if they exist, and put mermaid blocks instead.
    pattern = re.escape(title) + r'\s+!\[.*?\]\(.*?\)\s+'
    if re.search(pattern, readme_text):
        readme_text = re.sub(pattern, merm_code, readme_text)

with open(readme_path, 'w', encoding='utf-8') as f:
    f.write(readme_text)

print("Mermaid injected into Markdown.")

# Update the word document part
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level, style_name in [(1, 'Heading 1'), (2, 'Heading 2'), (3, 'Heading 3'), (4, 'Heading 4')]:
    h = doc.styles[style_name]
    h.font.name = 'Microsoft YaHei'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

lines = readme_text.split('\\n')
i = 0
in_mermaid = False
mermaid_buf = []

# Because we could not render mermaid to images in python easily without internet or heavy dependencies (playwright/puppeteer),
# we put the text block representation in the docx with a clear notice
while i < len(lines):
    line = lines[i].strip()
    
    if "```mermaid" in line:
        in_mermaid = True
        i += 1
        continue
        
    if in_mermaid:
        if line.startswith("```"):
            in_mermaid = False
            p = doc.add_paragraph("【注：由于系统网络原因，以上图表无法直接嵌入渲染，请通过支持 Mermaid 插件的 Markdown 阅读器（如 VSCode）查阅配套的 README.md 里的渲染高清图谱。】", style='Normal')
            p.runs[0].font.color.rgb = RGBColor(180, 0, 0)
            mermaid_buf = []
        else:
            mermaid_buf.append(line)
        i += 1
        continue

    # Skip regular code blocks
    if '```bash' in line or '```python' in line:
        i+=1
        while i < len(lines) and '```' not in lines[i]:
            i+=1
        i+=1
        continue
    
    if not line:
        i += 1
        continue
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        try:
            for run in p.runs:
                run.font.italic = True
        except:
            pass
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif len(line) > 0 and line[0].isdigit() and '. ' in line[:4]:
        idx = line.find('. ')
        doc.add_paragraph(line[idx+2:], style='List Number')
    elif line == '---':
        pass
    else:
        clean_line = line.replace('**', '').replace('`', '')
        doc.add_paragraph(clean_line)
        
    i += 1

doc_path = '职位平台-系统设计与评估报告(无插图附注版).docx'
doc.save(doc_path)
print(f"Word doc generated at {doc_path}")
