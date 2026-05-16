import os
import time
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Define the 6 diagrams
diagrams = {
    '顶层用例模型': '''graph LR
    subgraph 用户角色
        S[求职者 Student]
        C[企业人事 HR]
        A[系统管理员 Admin]
    end
    subgraph 职引未来 - 高校毕业生就业服务平台
        S -->|检索职位| UC1(职位搜索与筛选)
        S -->|查看分析| UC2(市场数据洞察查看)
        S -->|上传云端| UC3(简历云端管理解析)
        S -->|申请登记| UC4(投递追踪)
        S -->|在线沟通| UC5(即时通讯联络)
        
        C -->|发布及维保| UC6(职位发布管理)
        C -->|审核申请单| UC7(简历申请表审核)
        C -->|发送消息| UC8(发起面谈沟通)
        
        A -->|资质审核| UC9(企业资质全局管控)
        A -->|封禁违规| UC10(涉敏账号治理)
        A -->|维保监视| UC11(系统监控审核)
    end''',
    
    '系统宏观部署架构图': '''graph TD
    User((Web Users)) -->|HTTP/HTTPS| Nginx[Nginx 反代服务器]
    Nginx -->|WSGI/Gunicorn| Flask[Flask 应用服务集群]
    
    subgraph 核心服务层
        Flask -->|ORM 封装| SQLAlchemy(SQLAlchemy 引擎)
        Flask -->|读取/渲染| Visualizer(数据可视化模块)
        Visualizer -->|离线运算| Pandas(Pandas 引擎)
    end
    
    subgraph 存储与组件层
        SQLAlchemy -->|读写| MySQL[(MySQL 8.0 核心库)]
        Flask -->|读写| LocalStorage((本地对象存储 \n/uploads))
    end
    
    Pandas -.->|批量清洗入库| MySQL''',

    '核心功能模块结构图': '''graph TD
    Core[职引未来 - 核心业务中心] --> U(用户鉴权子系统)
    Core --> J(招聘投递子系统)
    Core --> D(数据洞察分析引擎)
    Core --> S(安全与风控模块)
    
    U --> UR(权限分配 - Student/Company/Admin)
    U --> UP(身份令牌分发与回话维系)
    
    J --> JS(多维搜索支持)
    J --> JA(状态机流转控制)
    J --> JC(求职双方通讯 Chat)
    
    D --> DA(薪酬地域 GroupBy)
    D --> DW(词云频率矩阵提取)
    
    S --> SL(表单及文件类型校验)
    S --> SS(黑名单阻断拦截)''',

    '类模型图 (Class Diagram)': '''classDiagram
    direction TB
    class User {
        +int id
        +string username
        +string pwd_hash
        +string role
    }
    class StudentProfile {
        +int id
        +string real_name
        +string university
    }
    class Resume {
        +int id
        +string file_path
    }
    class Company {
        +int id
        +string company_name
    }
    class Job {
        +int id
        +string title
        +int salary_min
    }
    User "1" *-- "1" StudentProfile : 关联附加
    StudentProfile "1" *-- "1..*" Resume : 拥有
    User "1" *-- "0..*" Company : 法权分配
    Company "1" *-- "1..*" Job : 分发管理''',
    
    '业务领域动态模型': '''sequenceDiagram
    participant S as 求职前端
    participant B as Flask 后端
    participant DB as MySQL 数据库
    participant C as HR 后台

    S->>B: POST /apply 发起申请脉冲
    B->>DB: 存库前确认及挂载限制
    DB-->>B: 回传验证许可
    B->>DB: 打入 Application挂靠记录(Pending)
    B-->>S: 告知 200 OK 投递已完成

    C->>B: GET /applications 请求待办
    B->>DB: 关联表调阅最新求职信源
    DB-->>B: 传回聚合实体
    B-->>C: 推送 JSON 映射待处列
    
    C->>B: PUT 审核邀约状态变更令
    B->>DB: 改写表值、重定义通讯解锁线
    B-->>C: 前端准许激活通讯窗口''',

    '数据库 ER 模型图': '''erDiagram
    USER ||--o| STUDENT_PROFILE : 挂载一对一
    USER ||--o| COMPANY : 公司全权
    COMPANY ||--o{ JOB : 下辖部署
    STUDENT_PROFILE ||--o{ RESUME : 创建文件组
    JOB ||--o{ APPLICATION : 被受发
    RESUME ||--o{ APPLICATION : 多路引源
    
    USER { int id PK } 
    JOB { int id PK } 
    APPLICATION { int id PK }'''
}

html_template = """
<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Mermaid Render</title>
  <script src="mermaid.min.js"></script>
  <script>
    mermaid.initialize({ startOnLoad: true, theme: 'default', htmlLabels: false });
  </script>
</head>
<body style="background: white; margin: 0; padding: 20px;">
  <div class="mermaid" id="diagram">
    _CODE_
  </div>
</body>
</html>
"""

os.makedirs('images', exist_ok=True)

# Generate screenshots
def generate_images():
    with sync_playwright() as p:
        browser = p.chromium.launch(channel="msedge", headless=True)
        page = browser.new_page()

        for name, code in diagrams.items():
            print(f"Rendering: {name}")
            html_content = html_template.replace('_CODE_', code)
            with open('temp.html', 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            page.goto(f"file://{os.path.abspath('temp.html')}")
            time.sleep(1)  # wait for mermaid to format DOM
            element = page.locator('#diagram')
            element.screenshot(path=f"images/{name}.png")
            print(f"Saved: images/{name}.png")

        browser.close()
    if os.path.exists('temp.html'):
        os.remove('temp.html')

def build_word_doc():
    doc = Document()
    
    # Enable support for Chinese font styles
    doc.styles['Normal'].font.name = u"微软雅黑"
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u"微软雅黑")

    doc.add_heading('职位平台-系统设计与评估报告', 0)
    
    # Copy all content from original word without the red highlights
    original = Document(r'c:\Users\床垫子\Desktop\积分单元\职位平台-系统设计与评估报告(无插图附注版).docx')
    for doc_p in original.paragraphs:
        if "【网络生成失败图片替代占位符】" in doc_p.text:
            continue
        
        # Replace headers locally with our image
        p = doc.add_paragraph(doc_p.text)
        p.style = doc_p.style

        for name in diagrams.keys():
            if name in doc_p.text and doc_p.style.name.startswith('Heading'):
                print(f"Inserting image for {name} into document")
                try:
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(f"images/{name}.png", width=Inches(5.0))
                except Exception as e:
                    print(f"Failed to insert {name}.png: {e}")

    save_path = '职位平台-系统设计与评估报告.docx'
    doc.save(save_path)
    print(f"Successfully generated final Word document at {save_path}")

if __name__ == "__main__":
    generate_images()
    build_word_doc()
