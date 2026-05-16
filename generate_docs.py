import os
import base64
import zlib
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 1. Define Mermaid Diagrams
diagrams = {
    'usecase': '''
graph LR
    subgraph 用户角色
        S[求职者 Student]
        C[企业人事 HR]
        A[系统管理员 Admin]
    end
    subgraph 职引未来 - 高校毕业生就业服务平台
        S -->|检索职位| UC1(职位搜索与筛选)
        S -->|查看分析| UC2(市场数据洞察查看)
        S -->|上传/更新| UC3(简历云端管理)
        S -->|申请登记| UC4(职位投递追踪)
        S -->|在线沟通| UC5(企业即时通讯)
        
        C -->|创建/下架| UC6(职位发布与维护)
        C -->|状态流转| UC7(简历申请审核)
        C -->|发送消息| UC8(发起面谈沟通)
        
        A -->|资质审核| UC9(企业资质管理)
        A -->|封禁/解禁| UC10(全局用户治理)
        A -->|监控| UC11(平台系统日志管理)
    end
    ''',
    'architecture': '''
graph TD
    subgraph 表现层 Views / Frontend
        UI[Jinja2 + Bootstrap UI]
        JS[原生 Fetch API & DOM 交互]
        ECH[ECharts & Pyecharts 可视化容器]
    end
    subgraph 业务逻辑层 Controllers / App
        AUTH[认证及鉴权中间件 \n @require_role]
        API[Search / Chat REST APIs]
        VIS[数据预计算大屏引擎 \n visualizer.py]
    end
    subgraph 数据持久层 Models / Database
        ORM[SQLAlchemy ORM 映射]
        DB[(MySQL 8.0 主数据库)]
        DATA[Pandas 离线清洗挂载 \n data_cleaner.py]
    end
    UI -->|HTTP GET/POST / 表单数据| AUTH
    JS -->|Fetch 异步请求| API
    ECH -->|静态文件加载| VIS
    AUTH --> ORM
    API --> ORM
    VIS -->|聚合分析| DATA
    DATA -->|初始导入| DB
    ORM -->|CRUD| DB
    ''',
    'flowchart': '''
flowchart TD
    A([求职者浏览岗位]) --> B{是否已登录?}
    B -- 否 --> C[跳转至登录/注册界面]
    C --> B
    B -- 是 --> D{是否已完善简历?}
    D -- 否 --> E[跳转至简历配置中心上传解析]
    E --> D
    D -- 是 --> F[发起简历投递请求]
    F --> G[系统校验企业接收容量/频率]
    G --> H[Application 表生成溯源快照]
    H --> I([企业端收到通知])
    I --> J{HR 资质审核}
    J -- 拒绝 --> K[更新状态为 Refused]
    J -- 邀约 --> L[开通即时通讯入口 Chat\n发出邀约]
    L --> M([双方在线沟通])
    ''',
    'class_diagram': '''
classDiagram
    class User {
        +int id
        +string email
        +string password_hash
        +string role(student/hr/admin)
        +verify_password(pwd)
    }
    class StudentProfile {
        +int user_id
        +string real_name
        +string university
        +string graduation_year
        +int target_salary
    }
    class Resume {
        +int id
        +int student_id
        +string file_path
        +string parsed_content
    }
    class Company {
        +int id
        +int owner_id
        +string company_name
        +string industry
        +bool is_verified
    }
    class Job {
        +int id
        +int company_id
        +string title
        +string location
        +int salary_min
        +int salary_max
    }
    User "1" *-- "1" StudentProfile : has
    User "1" *-- "1" Company : owns (if HR)
    StudentProfile "1" *-- "1..N" Resume : provides
    Company "1" *-- "1..N" Job : posts
    ''',
    'sequence_diagram': '''
sequenceDiagram
    participant S as 求职者 Student
    participant F as 前端界面 UI
    participant B as Flask 后端 API
    participant DB as MySQL 数据库
    participant C as 企业端 HR

    S->>F: 筛选并点击"投递该职位"
    F->>B: POST /apply (job_id, resume_id)
    B->>B: 鉴权安全网检验
    B->>DB: 查询任务与简历关系记录
    DB-->>B: 验证核对成功
    B->>DB: 创建 Application(状态=待处理)
    DB-->>B: 事务写成功
    B-->>F: 返回状态 200 OK
    F-->>S: 提示"投递成功"
    
    C->>F: 刷新后台面试处理栈
    F->>B: GET /applications?status=pending
    B->>DB: 查询当前 HR 从属的企业记录
    DB-->>B: 返回联表简历阵列
    B-->>F: JSON 返回申请人
    C->>F: 点击"发出面试邀约"
    F->>B: PUT /application/status
    B->>DB: 变更 Application, 解锁实时聊天通讯
    B-->>C: 切换通讯通道就续
    ''',
    'er_diagram': '''
erDiagram
    USER ||--o| STUDENT_PROFILE : "附加配置资料"
    USER ||--o| COMPANY : "关联组织账号(HR)"
    COMPANY ||--o{ JOB : "创建/下架"
    STUDENT_PROFILE ||--o{ RESUME : "拥有投递简历"
    JOB ||--o{ APPLICATION : "包含申请流转"
    RESUME ||--o{ APPLICATION : "附件快照快写"
    USER ||--o{ CHAT_LOG : "双向收发"
    
    USER {
        int id PK
        string email 
        string hashed_pwd
        string role
    }
    JOB {
        int id PK
        int company_id FK
        string title
        int salary
    }
    APPLICATION {
        int id PK
        int job_id FK
        int resume_id FK
        string status "pending/rejected/invited"
    }
    '''
}

def generate_kroki_url(diagram_text, dia_type="mermaid"):
    data = diagram_text.encode('utf-8')
    compressed = zlib.compress(data)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/{dia_type}/png/{encoded}"

os.makedirs('static/img', exist_ok=True)
images_paths = {}

print("Downloading generated diagrams from Kroki API...")
for name, code in diagrams.items():
    url = generate_kroki_url(code)
    try:
        resp = requests.get(url, timeout=10)
        img_path = f"static/img/{name}.png"
        with open(img_path, 'wb') as f:
            f.write(resp.content)
        images_paths[name] = img_path
        print(f"  Saved {img_path}")
    except Exception as e:
        print(f"Error fetching {name}: {e}")

# Inject into README.md
print("Updating README.md with diagram references...")
with open('README.md', 'r', encoding='utf-8') as f:
    readme_text = f.read()

readme_text = readme_text.replace("### 1.5 初始顶层用例模型", 
f"### 1.5 初始顶层用例模型\n\n![顶层用例模型]({images_paths.get('usecase', '')})\n\n")

readme_text = readme_text.replace("### 3.1 架构设计决策与技术选型依据", 
f"### 3.1 架构设计决策与技术选型依据\n\n#### 系统架构设计\n![系统架构设计图]({images_paths.get('architecture', '')})\n\n")

readme_text = readme_text.replace("### 4.1 问题域核心业务编码", 
f"### 4.1 问题域核心业务编码\n\n#### 系统核心投递与沟通流程图\n![系统流程图]({images_paths.get('flowchart', '')})\n\n")

readme_text = readme_text.replace("### 3.2 实体关系与数据库模型架构设计", 
f"### 3.2 实体关系与数据库模型架构设计\n\n#### 数据库 ER 模型图\n![数据模型ER图]({images_paths.get('er_diagram', '')})\n\n#### 业务领域静态模型（核心实体类图）\n![业务领域静态模型]({images_paths.get('class_diagram', '')})\n\n")

readme_text = readme_text.replace("### 3.3 核心拓展模块与可行性设计凭证", 
f"### 3.3 核心拓展模块与可行性设计凭证\n\n#### 业务领域动态模型（简历投递与邀约状态流转时序图）\n![业务涉及动态模型]({images_paths.get('sequence_diagram', '')})\n\n")

with open('README.md', 'w', encoding='utf-8') as f:
    f.write(readme_text)
print("README.md updated.")

# Convert to Word Document
print("Generating Word Document...")
doc = Document()

# Define Styles for Chinese Font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading1 = doc.styles['Heading 1']
heading1.font.name = 'Microsoft YaHei'
heading1.font.size = Pt(16)
heading1.font.color.rgb = RGBColor(0, 0, 0)
heading1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading2 = doc.styles['Heading 2']
heading2.font.name = 'Microsoft YaHei'
heading2.font.size = Pt(14)
heading2.font.color.rgb = RGBColor(0, 0, 0)
heading2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

heading3 = doc.styles['Heading 3']
heading3.font.name = 'Microsoft YaHei'
heading3.font.size = Pt(12)
heading3.font.color.rgb = RGBColor(0, 0, 0)
heading3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Parse Markdown to docx roughly
lines = readme_text.split('\\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    
    # Process Images
    if line.startswith('![') and '](' in line:
        img_path = line.split('](')[1].split(')')[0]
        if os.path.exists(img_path):
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            rt = doc.paragraphs[-1].add_run()
            try:
                rt.add_picture(img_path, width=Inches(5.0))
            except Exception as e:
                print(f"Could not add picture: {e}")
            
            # Subtitle
            lbl = line.split('![')[1].split(']')[0]
            p = doc.add_paragraph(lbl)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.style.font.size = Pt(10)
            except:
                pass
    elif line.startswith('# '):
        doc.add_heading(line[2:], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        try:
            for run in p.runs:
                run.font.italic = True
        except:
            pass
    elif line.startswith('- ') or line.startswith('* '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif len(line) > 0 and line[0].isdigit() and '. ' in line[:4]:
        idx = line.find('. ')
        doc.add_paragraph(line[idx+2:], style='List Number')
    elif line == '---' or line.startswith('```'):
        pass
    else:
        clean_line = line.replace('**', '').replace('`', '')
        doc.add_paragraph(clean_line)
        
    i += 1

doc_path = '职位平台-系统设计与评估报告.docx'
doc.save(doc_path)
print(f"Successfully generated {doc_path}")

